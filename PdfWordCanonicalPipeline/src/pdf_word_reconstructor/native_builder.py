from __future__ import annotations

import re
import shutil
import subprocess
import tempfile
import textwrap
import zipfile
from copy import deepcopy
from pathlib import Path
from statistics import median
from typing import Any, Iterable

from rapidfuzz import fuzz
from lxml import etree
from PIL import Image

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.part import Part
from docx.shared import Pt, RGBColor

from .common import normalize_text
from .docx_analyzer import iter_block_items


EMU_PER_PT = 12700
TWIPS_PER_PT = 20




SVG_NS = "http://schemas.microsoft.com/office/drawing/2016/SVG/main"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
SVG_EXT_URI = "{96DAC541-7B7A-43D3-8B79-37D633B846F1}"


def _attach_svg_to_picture(document_part, inline, svg_path: Path) -> str | None:
    """Attach a native SVG relationship to python-docx's PNG fallback.

    Word keeps the raster fallback in a:blip/r:embed and reads the SVG from an
    asvg:svgBlip extension.  This preserves vector fidelity while older Word
    versions can still display the fallback.
    """
    svg_path = Path(svg_path)
    if not svg_path.exists():
        return None
    package = document_part.package
    partname = package.next_partname("/word/media/image%d.svg")
    svg_part = Part(partname, "image/svg+xml", svg_path.read_bytes(), package)
    package.parts.append(svg_part)
    svg_rid = document_part.relate_to(svg_part, RT.IMAGE)
    blips = inline.xpath(".//*[local-name()='blip']")
    if not blips:
        return None
    blip = blips[0]
    ext_lst = next((child for child in blip if child.tag == f"{{{A_NS}}}extLst"), None)
    if ext_lst is None:
        ext_lst = etree.Element(f"{{{A_NS}}}extLst")
        blip.append(ext_lst)
    ext = etree.SubElement(ext_lst, f"{{{A_NS}}}ext")
    ext.set("uri", SVG_EXT_URI)
    svg_blip = etree.SubElement(ext, f"{{{SVG_NS}}}svgBlip", nsmap={"asvg": SVG_NS})
    svg_blip.set(f"{{{R_NS}}}embed", svg_rid)
    return svg_rid


def _add_picture_with_optional_svg(run, raster_path: Path, svg_path: Path | None, *, width: Pt, height: Pt):
    temp_raster: Path | None = None
    try:
        shape = run.add_picture(str(raster_path), width=width, height=height)
    except Exception:
        temp_raster = _rewrite_raster_as_docx_png(Path(raster_path))
        shape = run.add_picture(str(temp_raster), width=width, height=height)
    finally:
        if temp_raster is not None:
            try:
                temp_raster.unlink()
            except OSError:
                pass
    if svg_path:
        _attach_svg_to_picture(run.part, shape._inline, Path(svg_path))
    return shape


def _rewrite_raster_as_docx_png(raster_path: Path) -> Path:
    """Rewrite browser/Pillow-readable images that python-docx rejects.

    Some Mathpix JPEGs are valid enough for browsers and Pillow but have headers
    that python-docx refuses.  The reconstructed DOCX should not fail the entire
    page for that, so normalize the bytes to a fresh PNG fallback.
    """
    handle = tempfile.NamedTemporaryFile(prefix="bw_docx_image_", suffix=".png", delete=False)
    target = Path(handle.name)
    handle.close()
    try:
        with Image.open(raster_path) as image:
            image.load()
            if image.mode not in {"RGB", "RGBA", "L"}:
                image = image.convert("RGBA" if "A" in image.getbands() else "RGB")
            image.save(target, format="PNG")
        return target
    except Exception as exc:
        try:
            target.unlink()
        except OSError:
            pass
        raise RuntimeError(f"Unsupported image asset for DOCX: {raster_path}") from exc

def _append_latex_as_omml(target_paragraph, latex: str, diagnostics: dict[str, Any] | None = None) -> int:
    """Convert one trusted Mathpix LaTeX donor to native OMML through Pandoc."""
    latex = str(latex or "").strip()
    if not latex:
        if diagnostics is not None:
            diagnostics["latex_omml_failure"] = {"reason": "empty-latex"}
        return 0
    if shutil.which("pandoc") is None:
        if diagnostics is not None:
            diagnostics["latex_omml_failure"] = {"reason": "pandoc-not-found"}
        return 0
    with tempfile.TemporaryDirectory(prefix="bw_latex_omml_") as td:
        td_path = Path(td)
        md = td_path / "equation.md"
        docx = td_path / "equation.docx"
        md.write_text("$$\n" + latex + "\n$$\n", encoding="utf-8")
        try:
            completed = subprocess.run(["pandoc", str(md), "-o", str(docx)], capture_output=True, text=True, timeout=30)
        except (OSError, subprocess.SubprocessError) as exc:
            if diagnostics is not None:
                diagnostics["latex_omml_failure"] = {"reason": "pandoc-exception", "error": str(exc)}
            return 0
        if completed.returncode != 0 or not docx.exists():
            if diagnostics is not None:
                diagnostics["latex_omml_failure"] = {
                    "reason": "pandoc-returned-no-docx",
                    "returncode": completed.returncode,
                    "stderr": (completed.stderr or "")[-1200:],
                }
            return 0
        with zipfile.ZipFile(docx) as archive:
            root = etree.fromstring(archive.read("word/document.xml"))
        maths = root.xpath(".//*[local-name()='oMath']")
        if not maths and diagnostics is not None:
            diagnostics["latex_omml_failure"] = {"reason": "pandoc-docx-without-omml"}
        for math in maths:
            target_paragraph._p.append(deepcopy(math))
        return len(maths)


def _clean_pdf_text(text: str) -> str:
    text = text or ""
    text = text.replace("\u00ad", "")
    text = re.sub(r"([\wά-ώΆ-Ώ])-\s*\n\s*([\wά-ώΆ-Ώ])", r"\1\2", text)
    text = re.sub(r"\s*\n\s*", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def _max_span_size(regions_by_id: dict[str, dict[str, Any]], region_ids: Iterable[str], fallback: float) -> float:
    values: list[float] = []
    for rid in region_ids:
        region = regions_by_id.get(rid, {})
        for line in region.get("lines", []):
            for span in line.get("spans", []):
                if str(span.get("text", "")).strip():
                    values.append(float(span.get("size_pt") or fallback))
    return max(values, default=fallback)


def _dominant_span_size(regions_by_id: dict[str, dict[str, Any]], region_ids: Iterable[str], fallback: float) -> float:
    weights: dict[float, float] = {}
    for rid in region_ids:
        region = regions_by_id.get(rid, {})
        for line in region.get("lines", []):
            for span in line.get("spans", []):
                text = str(span.get("text") or "")
                if not text.strip():
                    continue
                try:
                    raw_size = float(span.get("size_pt") or fallback)
                except (TypeError, ValueError):
                    raw_size = float(fallback)
                size = round(raw_size * 2.0) / 2.0
                weights[size] = weights.get(size, 0.0) + max(1.0, len(text.strip()))
    if not weights:
        return float(fallback)
    return max(weights.items(), key=lambda item: (item[1], item[0]))[0]


def _line_height(regions_by_id: dict[str, dict[str, Any]], region_ids: Iterable[str], fallback: float) -> float:
    pitches: list[float] = []
    heights: list[float] = []
    for rid in region_ids:
        region = regions_by_id.get(rid, {})
        lines = sorted(
            [
                line for line in region.get("lines", [])
                if isinstance(line.get("bbox"), (list, tuple)) and len(line.get("bbox")) == 4
            ],
            key=lambda line: (float(line["bbox"][1]), float(line["bbox"][0])),
        )
        for previous, current in zip(lines, lines[1:]):
            pitch = float(current["bbox"][1]) - float(previous["bbox"][1])
            if fallback * 0.65 <= pitch <= fallback * 2.15:
                pitches.append(pitch)
        for line in lines:
            box = line.get("bbox", [0, 0, 0, 0])
            height = float(box[3]) - float(box[1])
            if height > 2:
                heights.append(height)
    if pitches:
        value = median(pitches)
        value = max(fallback * 1.0, min(fallback * 1.9, value))
        return round(value * 2.0) / 2.0
    if not heights:
        return round((fallback * 1.15) * 2.0) / 2.0
    value = max(fallback * 1.08, min(fallback * 1.55, median(heights)))
    return round(value * 2.0) / 2.0


def _pdf_line_records(regions_by_id: dict[str, dict[str, Any]], region_ids: Iterable[str]) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    seen: set[tuple] = set()
    for rid in region_ids:
        region = regions_by_id.get(rid, {})
        for line in region.get("lines", []):
            box = line.get("bbox")
            if not isinstance(box, (list, tuple)) or len(box) != 4:
                continue
            text = "".join(str(span.get("text") or "") for span in line.get("spans", []))
            if not text.strip():
                continue
            key = (*[round(float(value), 1) for value in box], normalize_text(text)[:80])
            if key in seen:
                continue
            seen.add(key)
            records.append({
                "bbox": [float(value) for value in box],
                "text": text,
            })
    return sorted(records, key=lambda row: (row["bbox"][1], row["bbox"][0]))


def _pdf_text_geometry(regions_by_id: dict[str, dict[str, Any]], region_ids: Iterable[str]) -> dict[str, Any] | None:
    lines = _pdf_line_records(regions_by_id, region_ids)
    if not lines:
        boxes = [
            region.get("bbox")
            for rid in region_ids
            for region in [regions_by_id.get(rid, {})]
            if isinstance(region.get("bbox"), (list, tuple)) and len(region.get("bbox")) == 4
        ]
        if not boxes:
            return None
        x0 = min(float(box[0]) for box in boxes)
        y0 = min(float(box[1]) for box in boxes)
        x1 = max(float(box[2]) for box in boxes)
        y1 = max(float(box[3]) for box in boxes)
        return {
            "line_count": 0,
            "bbox": [x0, y0, x1, y1],
            "height_pt": y1 - y0,
            "width_pt": x1 - x0,
            "line_pitch_pt": None,
        }
    x0 = min(float(line["bbox"][0]) for line in lines)
    y0 = min(float(line["bbox"][1]) for line in lines)
    x1 = max(float(line["bbox"][2]) for line in lines)
    y1 = max(float(line["bbox"][3]) for line in lines)
    pitches: list[float] = []
    for previous, current in zip(lines, lines[1:]):
        pitch = float(current["bbox"][1]) - float(previous["bbox"][1])
        if 2.0 <= pitch <= 40.0:
            pitches.append(pitch)
    return {
        "line_count": len(lines),
        "bbox": [x0, y0, x1, y1],
        "height_pt": y1 - y0,
        "width_pt": x1 - x0,
        "line_pitch_pt": median(pitches) if pitches else None,
    }


def _flow_geometry_fit(
    page_regions: dict[str, dict[str, Any]],
    item: dict[str, Any],
    *,
    column_width_pt: float,
    font_size_pt: float,
    line_height_pt: float,
    gap_pt: float,
) -> dict[str, Any] | None:
    geometry = _pdf_text_geometry(page_regions, item.get("region_ids", []))
    if not geometry:
        return None
    line_count = int(geometry.get("line_count") or 0)
    if line_count <= 0:
        return None
    pdf_height = float(geometry.get("height_pt") or 0.0)
    pdf_width = float(geometry.get("width_pt") or 0.0)
    line_height = float(line_height_pt or 0.0)
    font_size = float(font_size_pt or 0.0)
    word_text_height = max(font_size, ((line_count - 1) * line_height) + font_size)
    word_flow_advance = line_count * line_height
    end_delta = word_text_height - pdf_height
    line_pitch = geometry.get("line_pitch_pt")
    pitch_delta = None if line_pitch is None else line_height - float(line_pitch)
    if abs(end_delta) <= 12.0:
        status = "ok"
    elif abs(end_delta) <= 24.0:
        status = "review"
    else:
        status = "bad"
    box = geometry["bbox"]
    return {
        "status": status,
        "pdf_y0_pt": round(float(box[1]), 2),
        "pdf_y1_pt": round(float(box[3]), 2),
        "estimated_word_y1_pt": round(float(box[1]) + word_text_height, 2),
        "pdf_line_count": line_count,
        "pdf_text_height_pt": round(pdf_height, 2),
        "pdf_text_width_pt": round(pdf_width, 2),
        "word_column_width_pt": round(float(column_width_pt), 2),
        "column_width_delta_pt": round(float(column_width_pt) - pdf_width, 2),
        "pdf_line_pitch_pt": round(float(line_pitch), 2) if line_pitch is not None else None,
        "word_font_size_pt": round(font_size, 2),
        "word_line_height_pt": round(line_height, 2),
        "line_pitch_delta_pt": round(float(pitch_delta), 2) if pitch_delta is not None else None,
        "word_text_height_pt": round(word_text_height, 2),
        "word_flow_advance_pt": round(word_flow_advance, 2),
        "text_end_delta_pt": round(end_delta, 2),
        "gap_pt": round(float(gap_pt), 2),
        "policy": "pdf-paragraph-geometry-fit-before-render",
    }


def _summarize_flow_geometry_fits(rows: list[dict[str, Any]]) -> dict[str, Any]:
    deltas = [float(row.get("text_end_delta_pt") or 0.0) for row in rows]
    pitch_deltas = [
        float(row["line_pitch_delta_pt"])
        for row in rows
        if row.get("line_pitch_delta_pt") is not None
    ]
    if not deltas:
        return {
            "policy": "pdf-paragraph-geometry-fit-before-render",
            "count": 0,
        }
    return {
        "policy": "pdf-paragraph-geometry-fit-before-render",
        "count": len(rows),
        "ok_count": sum(1 for row in rows if row.get("status") == "ok"),
        "review_count": sum(1 for row in rows if row.get("status") == "review"),
        "bad_count": sum(1 for row in rows if row.get("status") == "bad"),
        "average_abs_text_end_delta_pt": round(sum(abs(value) for value in deltas) / len(deltas), 2),
        "max_abs_text_end_delta_pt": round(max(abs(value) for value in deltas), 2),
        "average_line_pitch_delta_pt": (
            round(sum(pitch_deltas) / len(pitch_deltas), 2)
            if pitch_deltas else None
        ),
    }


def _set_cell_free_document_settings(doc: Document) -> None:
    settings = doc.settings._element
    for tag, value in (("w:autoHyphenation", "1"), ("w:doNotHyphenateCaps", "0")):
        element = settings.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            settings.append(element)
        element.set(qn("w:val"), value)


def _set_mirror_margins(doc: Document, enabled: bool) -> None:
    settings = doc.settings._element
    element = settings.find(qn("w:mirrorMargins"))
    if enabled and element is None:
        settings.append(OxmlElement("w:mirrorMargins"))
    elif not enabled and element is not None:
        settings.remove(element)


def _set_run_language(run, lang: str = "el-GR") -> None:
    r_pr = run._r.get_or_add_rPr()
    lang_el = r_pr.find(qn("w:lang"))
    if lang_el is None:
        lang_el = OxmlElement("w:lang")
        r_pr.append(lang_el)
    lang_el.set(qn("w:val"), lang)
    lang_el.set(qn("w:eastAsia"), lang)


def _set_paragraph_borders(paragraph, color: str = "C00000", fill: str | None = None, size: int = 8) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if fill:
        shd = p_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            p_pr.append(shd)
        shd.set(qn("w:fill"), fill)
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "1")
        node.set(qn("w:color"), color)


def _set_frame(paragraph, bbox: list[float]) -> None:
    x0, y0, x1, y1 = map(float, bbox)
    p_pr = paragraph._p.get_or_add_pPr()
    frame = p_pr.find(qn("w:framePr"))
    if frame is None:
        frame = OxmlElement("w:framePr")
        p_pr.insert(0, frame)
    frame.set(qn("w:w"), str(max(1, round((x1 - x0) * TWIPS_PER_PT))))
    frame.set(qn("w:h"), str(max(1, round((y1 - y0) * TWIPS_PER_PT))))
    frame.set(qn("w:x"), str(round(x0 * TWIPS_PER_PT)))
    frame.set(qn("w:y"), str(round(y0 * TWIPS_PER_PT)))
    frame.set(qn("w:hAnchor"), "page")
    frame.set(qn("w:vAnchor"), "page")
    frame.set(qn("w:wrap"), "around")
    frame.set(qn("w:hRule"), "exact")
    frame.set(qn("w:hSpace"), "0")
    frame.set(qn("w:vSpace"), "0")
    frame.set(qn("w:anchorLock"), "1")


def _inline_to_anchor(inline, x_pt: float, y_pt: float, wrap: str = "none", relative_height: int = 251658240):
    anchor = OxmlElement("wp:anchor")
    attributes = {
        "distT": "0",
        "distB": "0",
        "distL": "0",
        "distR": "0",
        "simplePos": "0",
        "relativeHeight": str(relative_height),
        "behindDoc": "0",
        "locked": "0",
        "layoutInCell": "1",
        "allowOverlap": "1",
    }
    for key, value in attributes.items():
        anchor.set(key, value)

    simple_pos = OxmlElement("wp:simplePos")
    simple_pos.set("x", "0")
    simple_pos.set("y", "0")
    anchor.append(simple_pos)

    position_h = OxmlElement("wp:positionH")
    position_h.set("relativeFrom", "page")
    pos_h = OxmlElement("wp:posOffset")
    pos_h.text = str(round(x_pt * EMU_PER_PT))
    position_h.append(pos_h)
    anchor.append(position_h)

    position_v = OxmlElement("wp:positionV")
    position_v.set("relativeFrom", "page")
    pos_v = OxmlElement("wp:posOffset")
    pos_v.text = str(round(y_pt * EMU_PER_PT))
    position_v.append(pos_v)
    anchor.append(position_v)

    children = {child.tag: child for child in inline}
    for name in ("wp:extent", "wp:effectExtent"):
        child = children.get(qn(name))
        if child is not None:
            anchor.append(deepcopy(child))

    if wrap == "square":
        wrap_node = OxmlElement("wp:wrapSquare")
        wrap_node.set("wrapText", "bothSides")
    else:
        wrap_node = OxmlElement("wp:wrapNone")
    anchor.append(wrap_node)

    for name in ("wp:docPr", "wp:cNvGraphicFramePr", "a:graphic"):
        child = children.get(qn(name))
        if child is not None:
            anchor.append(deepcopy(child))

    inline.getparent().replace(inline, anchor)
    return anchor


def _add_floating_picture(doc: Document, image_path: Path, bbox: list[float], wrap: str = "none", z: int = 251658240, svg_path: Path | None = None) -> None:
    x0, y0, x1, y1 = map(float, bbox)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(1)
    run = p.add_run()
    run.font.size = Pt(1)
    inline_shape = _add_picture_with_optional_svg(
        run, image_path, svg_path,
        width=Pt(max(1, x1 - x0)), height=Pt(max(1, y1 - y0)),
    )
    inline = inline_shape._inline
    _inline_to_anchor(inline, x0, y0, wrap=wrap, relative_height=z)


def _set_paragraph_default_font(paragraph: Paragraph, font_size: float, font_name: str = "Times New Roman") -> None:
    """Set paragraph-mark run properties so inline OMML inherits the frame scale."""
    ppr = paragraph._p.get_or_add_pPr()
    rpr = ppr.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        ppr.append(rpr)
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attr}"), font_name)
    half_points = str(max(10, round(float(font_size) * 2)))
    for tag in ("w:sz", "w:szCs"):
        node = rpr.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            rpr.append(node)
        node.set(qn("w:val"), half_points)


def _estimated_callout_lines(text: str, width_pt: float, font_size: float) -> int:
    # Times New Roman prose averages roughly half an em per character. A slightly
    # conservative estimate is safer than allowing a Word frame to grow into the
    # next equation or sidebar object.
    usable_width = max(16.0, width_pt - 8.0)
    chars_per_line = max(8, int(usable_width / max(3.0, font_size * 0.50)))
    total = 0
    for raw_line in (text or "").splitlines() or [""]:
        line = re.sub(r"\s+", " ", raw_line).strip()
        if not line:
            total += 1
            continue
        total += max(1, len(textwrap.wrap(line, width=chars_per_line, break_long_words=True, break_on_hyphens=True)))
    return max(1, total)


def _fit_callout_font(
    callout: dict[str, Any],
    text: str,
    body_size: float,
    native_math_count: int = 0,
) -> tuple[float, float, int]:
    x0, y0, x1, y1 = map(float, callout["bbox"])
    width, height = max(1.0, x1 - x0), max(1.0, y1 - y0)
    source_size = float(callout.get("semantic", {}).get("stats", {}).get("weighted_size") or 0.0)
    preferred = source_size if 6.0 <= source_size <= 10.5 else body_size * 0.82
    preferred = max(6.0, min(8.9, preferred))
    best = 5.8
    best_lines = _estimated_callout_lines(text, width, best)
    # Native fractions/superscripts need more vertical room than an ordinary
    # line. Reserve about 1.8 prose lines for each absorbed OMML expression.
    math_line_equivalent = max(0, int(native_math_count)) * 1.8
    candidate = preferred
    while candidate >= 5.8:
        lines = _estimated_callout_lines(text, width, candidate)
        line_height = max(candidate * 1.04, candidate + 0.25)
        required = (lines + math_line_equivalent) * line_height + 6.5
        if required <= height:
            best, best_lines = candidate, lines
            break
        candidate -= 0.1
    line_height = max(best * 1.04, best + 0.25)
    return round(best, 2), round(line_height, 2), best_lines


def _add_callout(
    doc: Document,
    callout: dict[str, Any],
    text: str,
    body_size: float,
    source_paragraphs: dict[str, Paragraph] | None = None,
    paragraph_ids: list[str] | None = None,
) -> dict[str, Any]:
    p = doc.add_paragraph()
    _set_frame(p, callout["bbox"])
    _set_paragraph_borders(p, color="FF0000", fill="FFFFFF", size=8)
    p.paragraph_format.left_indent = Pt(2.5)
    p.paragraph_format.right_indent = Pt(2.5)
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_together = True
    p.paragraph_format.widow_control = False
    expected_math_count = 0
    if source_paragraphs and paragraph_ids:
        expected_math_count = sum(
            len(source_paragraphs[pid]._p.xpath(".//m:oMath"))
            for pid in paragraph_ids if pid in source_paragraphs
        )
    font_size, line_height, estimated_lines = _fit_callout_font(
        callout, text, body_size, native_math_count=expected_math_count
    )
    _set_paragraph_default_font(p, font_size)
    p.paragraph_format.line_spacing_rule = (
        WD_LINE_SPACING.AT_LEAST if expected_math_count else WD_LINE_SPACING.EXACTLY
    )
    p.paragraph_format.line_spacing = Pt(line_height)
    native_info = {"copied_nodes": 0, "math_count": 0}
    if source_paragraphs and paragraph_ids:
        native_info = _append_source_paragraph_content(
            p, source_paragraphs, paragraph_ids, font_size,
            flatten_math_paragraphs=True,
        )
    else:
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(font_size)
        _set_run_language(run)
    return {
        "id": callout.get("id"),
        "font_size_pt": font_size,
        "line_height_pt": line_height,
        "estimated_lines": estimated_lines,
        "bbox": callout.get("bbox"),
        "source_paragraphs": list(paragraph_ids or []),
        "native_math_count": int(native_info.get("math_count", 0)),
        "contained_visual_groups": list(callout.get("contained_visual_groups", [])),
    }


def _clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)



def _source_paragraph_map(docx_path: Path) -> tuple[Document, dict[str, Paragraph]]:
    """Re-open the source DOCX and reproduce the analyzer's paragraph IDs.

    Keeping the live OOXML paragraph elements lets the builder preserve native
    Office Math (OMML) instead of replacing equations with raster crops.
    """
    source_doc = Document(docx_path)
    mapping: dict[str, Paragraph] = {}
    index = 0
    for item in iter_block_items(source_doc):
        if isinstance(item, Paragraph):
            mapping[f"d-p{index:05d}"] = item
            index += 1
        elif isinstance(item, Table):
            for row in item.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        mapping[f"d-p{index:05d}"] = paragraph
                        index += 1
    return source_doc, mapping


def _strip_visuals_from_clone(node) -> None:
    """Remove drawings from copied source paragraphs.

    Figures are positioned independently from the PDF page map. Copying a source
    drawing here would duplicate it, while OMML and ordinary text remain useful.
    """
    visual_tags = {qn("w:drawing"), qn("w:pict"), qn("w:object")}
    for descendant in list(node.iter()):
        if descendant.tag in visual_tags:
            parent = descendant.getparent()
            if parent is not None:
                parent.remove(descendant)


def _clone_has_visible_content(node) -> bool:
    visible_tags = {
        qn("w:t"), qn("w:tab"), qn("w:br"), qn("w:cr"),
        qn("m:oMath"), qn("m:oMathPara"), qn("m:t"),
    }
    return node.tag in visible_tags or any(desc.tag in visible_tags for desc in node.iter())


def _append_source_paragraph_content(
    target,
    source_paragraphs: dict[str, Paragraph],
    paragraph_ids: list[str],
    font_size: float,
    color: RGBColor | None = None,
    force_bold: bool = False,
    italic: bool = False,
    flatten_math_paragraphs: bool = False,
) -> dict[str, Any]:
    """Copy text + native OMML while discarding source layout and drawings."""
    copied = 0
    math_count = 0
    for pos, pid in enumerate(paragraph_ids):
        source = source_paragraphs.get(pid)
        if source is None:
            continue
        if copied and pos:
            target.add_run().add_break(WD_BREAK.LINE)
        for child in source._p:
            if child.tag in {
                qn("w:pPr"), qn("w:bookmarkStart"), qn("w:bookmarkEnd"),
                qn("w:proofErr"), qn("w:permStart"), qn("w:permEnd"),
            }:
                continue
            clone = deepcopy(child)
            _strip_visuals_from_clone(clone)
            if not _clone_has_visible_content(clone):
                continue
            math_count += sum(1 for element in clone.iter() if element.tag == qn("m:oMath"))
            if flatten_math_paragraphs and clone.tag == qn("m:oMathPara"):
                # A display-math paragraph inside a fixed-height Word frame is
                # prone to clipping. Keep the native OMML but place its oMath
                # child inline on the preceding explicit line break.
                maths = [deepcopy(child) for child in clone if child.tag == qn("m:oMath")]
                for math in maths:
                    target._p.append(math)
                    copied += 1
                continue
            target._p.append(clone)
            copied += 1

    # Apply the target page's typography to normal text runs. OMML remains native
    # and inherits the paragraph's surrounding font metrics in Word.
    for run in target.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(font_size)
        if force_bold:
            run.bold = True
        if italic:
            run.italic = True
        if color:
            run.font.color.rgb = color
        _set_run_language(run)
    return {"copied_nodes": copied, "math_count": math_count}


def _strong_match_docx_indexes(item: dict[str, Any], matches: dict[str, dict[str, Any]]) -> list[int]:
    indexes: list[int] = []
    for rid in item.get("region_ids", []):
        match = matches.get(rid)
        if match and match.get("status") in {"strong", "medium"}:
            indexes.extend(int(value) for value in match.get("docx_indexes", []))
    return sorted(set(indexes))


def _math_signature(text: str) -> str:
    text = (text or "").lower()
    replacements = {
        "−": "-", "–": "-", "—": "-", "⋅": "*", "·": "*", "×": "*",
        "ν": "v", "𝜈": "v", "λ": "l", "𝜆": "l", "ε": "e", "𝜀": "e",
        "α": "a", "β": "b", "γ": "g", "θ": "th", "μ": "m", "π": "p",
        "ρ": "r", "σ": "s", "τ": "t", "φ": "f", "χ": "x", "ω": "w",
        "∞": "inf", "→": "->", "𝛥": "d", "δ": "d", "Δ": "d",
        "𝐸": "e", "𝐽": "j", "ℎ": "h", "𝑛": "n", "𝑐": "c",
    }
    for a, b in replacements.items():
        text = text.replace(a, b)
    text = re.sub(r"[^a-z0-9+\-*/=<>]", "", text)
    return text


def _latex_math_signature(latex: str) -> str:
    text = (latex or "").lower()
    replacements = {
        "\\alpha": "a", "\\beta": "b", "\\gamma": "g", "\\delta": "d",
        "\\Delta": "d", "\\epsilon": "e", "\\varepsilon": "e",
        "\\theta": "th", "\\lambda": "l", "\\mu": "m", "\\nu": "v",
        "\\pi": "p", "\\rho": "r", "\\sigma": "s", "\\tau": "t",
        "\\phi": "f", "\\varphi": "f", "\\chi": "x", "\\omega": "w",
        "\\infty": "inf", "\\to": "->", "\\rightarrow": "->",
        "\\cdot": "*", "\\times": "*",
    }
    for source, target in replacements.items():
        text = text.replace(source.lower(), target)
    text = text.replace("\\mathrm", "")
    text = text.replace("\\text", "")
    text = text.replace("\\operatorname", "")
    text = text.replace("\\mathbf", "")
    text = text.replace("\\boldsymbol", "")
    text = text.replace("\\left", "")
    text = text.replace("\\right", "")
    previous = None
    while previous != text:
        previous = text
        text = re.sub(r"\\frac\s*\{([^{}]*)\}\s*\{([^{}]*)\}", r"\1/\2", text)
        text = re.sub(r"\\(?:mathrm|text|operatorname)\s*\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\begin\{[^{}]+\}|\\end\{[^{}]+\}", "", text)
    text = re.sub(r"\\[a-z]+", "", text)
    text = text.replace("{", "").replace("}", "")
    return _math_signature(text)


def _item_math_signatures(item: dict[str, Any]) -> list[tuple[str, str]]:
    signatures: list[tuple[str, str]] = []
    markdown_sig = _latex_math_signature(str(item.get("markdown_latex") or ""))
    if markdown_sig:
        signatures.append(("markdown-latex", markdown_sig))
    pdf_sig = _math_signature(str(item.get("text", "")))
    if pdf_sig:
        signatures.append(("pdf-text", pdf_sig))
    return signatures


def _math_signature_score(left: str, right: str) -> tuple[float, float]:
    if not left or not right:
        return 0.0, 0.0
    similarity = float(fuzz.WRatio(left, right))
    score = similarity
    if left in right or right in left:
        coverage = min(len(left), len(right)) / max(len(left), len(right))
        score = max(score, 82.0 + 18.0 * coverage)
    return score, similarity


def _native_math_candidate_indexes(
    docx_analysis: dict[str, Any],
    range_start: int,
    range_end: int,
    docx_donor_map: dict[str, Any] | None = None,
) -> list[int]:
    if docx_donor_map is not None:
        indexes: list[int] = []
        for candidate in docx_donor_map.get("mathCandidates", []) or []:
            try:
                index = int(candidate.get("index"))
            except Exception:
                continue
            if int(range_start) <= index <= int(range_end):
                indexes.append(index)
        return sorted(set(indexes))
    candidates: list[int] = []
    for paragraph in docx_analysis.get("paragraphs", []):
        index = int(paragraph["index"])
        if index < int(range_start) or index > int(range_end):
            continue
        if int(paragraph.get("omml_count", 0)) <= 0:
            continue
        prose_len = len(normalize_text(paragraph.get("text", "")))
        if paragraph.get("is_math_only") or prose_len <= 24:
            candidates.append(index)
    return candidates


def _best_math_candidate(
    pool: list[int],
    records: dict[int, dict[str, Any]],
    item_signatures: list[tuple[str, str]],
    midpoint: float,
    *,
    wide_markdown_only: bool = False,
) -> tuple[float, int, float, str, str] | None:
    best: tuple[float, int, float, str, str] | None = None
    for index in pool:
        record = records[index]
        omml_sig = _math_signature(str(record.get("omml_text", "")))
        source = ""
        signature = ""
        similarity = 0.0
        base_score = 0.0
        for candidate_source, candidate_sig in item_signatures:
            if wide_markdown_only and candidate_source != "markdown-latex":
                continue
            candidate_score, candidate_similarity = _math_signature_score(candidate_sig, omml_sig)
            if candidate_score > base_score:
                base_score = candidate_score
                similarity = candidate_similarity
                source = candidate_source
                signature = candidate_sig
        if not source:
            continue
        distance = abs(index - midpoint)
        proximity = max(0.0, 100.0 - distance * 7.0)
        score = base_score * 0.78 + proximity * 0.22
        if source == "markdown-latex" and base_score >= 82.0:
            score = max(score, base_score)
        if wide_markdown_only:
            min_score = 96.0 if len(signature) < 6 else 86.0
            if base_score < min_score:
                continue
            score = base_score - min(10.0, distance * 0.03)
            source = "markdown-latex-wide-docx-omml"
        if best is None or score > best[0]:
            best = (score, index, similarity, source, signature)
    return best


def _assign_native_math_sources(
    flow: list[dict[str, Any]],
    matches: dict[str, dict[str, Any]],
    docx_analysis: dict[str, Any],
    candidate_range: list[int] | tuple[int, int] | None = None,
    used_native_math_indexes: set[int] | None = None,
    docx_donor_map: dict[str, Any] | None = None,
) -> None:
    """Pair displayed PDF equations with nearby OMML by content and position.

    The v0.6 wide run exposed a dangerous failure mode: when neighbouring prose was
    weakly aligned, the old positional picker could insert a perfectly valid but
    completely unrelated formula.  v0.7 uses the OMML text itself as the primary
    signal and falls back to the raster crop when confidence is low.
    """
    records = {int(p["index"]): p for p in docx_analysis.get("paragraphs", [])}
    range_start, range_end = (candidate_range or [0, 10**9])[:2]
    candidates = _native_math_candidate_indexes(docx_analysis, int(range_start), int(range_end), docx_donor_map)
    wide_start = max(0, int(range_start) - 140)
    wide_end = int(range_end) + 140
    wide_candidates = _native_math_candidate_indexes(docx_analysis, wide_start, wide_end, docx_donor_map)
    used: set[int] = used_native_math_indexes if used_native_math_indexes is not None else set()
    hints = [_strong_match_docx_indexes(item, matches) if item.get("type") == "text" else [] for item in flow]

    for pos, item in enumerate(flow):
        if item.get("type") != "visual" or item.get("semantic_type") != "equation":
            continue
        prev_end = None
        next_start = None
        for scan in range(pos - 1, -1, -1):
            if hints[scan]:
                prev_end = max(hints[scan])
                break
        for scan in range(pos + 1, len(flow)):
            if hints[scan]:
                next_start = min(hints[scan])
                break

        pool: list[int] = []
        for index in candidates:
            if index in used:
                continue
            if prev_end is not None and index <= prev_end:
                continue
            if next_start is not None and index >= next_start:
                continue
            pool.append(index)
        if not pool:
            anchor = prev_end if prev_end is not None else next_start
            if anchor is not None:
                pool = [index for index in candidates if index not in used and abs(index - anchor) <= 8]

        item_signatures = _item_math_signatures(item)
        if not item_signatures:
            continue
        if pool:
            midpoint = ((prev_end if prev_end is not None else pool[0]) + (next_start if next_start is not None else pool[-1])) / 2.0
            best = _best_math_candidate(pool, records, item_signatures, midpoint)
        else:
            anchors = [value for value in [prev_end, next_start] if value is not None]
            midpoint = sum(anchors) / len(anchors) if anchors else (int(range_start) + int(range_end)) / 2.0
            best = None
        if (best is None or best[0] < 58.0) and any(source == "markdown-latex" for source, _sig in item_signatures):
            expanded_pool = [index for index in wide_candidates if index not in used and index not in pool]
            wide_best = _best_math_candidate(
                expanded_pool, records, item_signatures, midpoint, wide_markdown_only=True
            )
            if wide_best is not None:
                best = wide_best
        if best is None or best[0] < 58.0:
            item["native_math_rejected"] = {
                "reason": "low-confidence",
                "item_signatures": [{"source": source, "signature": signature} for source, signature in item_signatures],
                "best_score": round(best[0], 2) if best else 0.0,
            }
            continue
        score, chosen, similarity, source, signature = best
        used.add(chosen)
        record = records[chosen]
        item["native_math_paragraph_ids"] = [record["id"]]
        item["native_math_docx_indexes"] = [chosen]
        item["native_math_match_score"] = round(score, 2)
        item["native_math_text_score"] = round(similarity, 2)
        item["native_math_signature_source"] = source
        item["native_math_item_signature"] = signature
        item["native_math_docx_signature"] = _math_signature(str(record.get("omml_text", "")))


def _callout_source_paragraph_ids(
    callout: dict[str, Any],
    match: dict[str, Any] | None,
    docx_analysis: dict[str, Any],
) -> list[str]:
    """Return the matched callout paragraphs plus adjacent OMML kept inside it.

    The PDF may split an equation visually contained by the callout into a second
    region. In the converted DOCX that equation is commonly the immediately
    following OMML-only paragraph. It belongs inside the framed callout, not in
    main page flow.
    """
    if not match or match.get("status") != "strong":
        return []
    ids = list(dict.fromkeys(match.get("docx_paragraphs", [])))
    if not callout.get("contained_visual_groups"):
        return ids
    indexes = [int(value) for value in match.get("docx_indexes", [])]
    if not indexes:
        return ids
    records = {int(p["index"]): p for p in docx_analysis.get("paragraphs", [])}
    cursor = max(indexes) + 1
    # One callout may contain more than one formula, but stop at the first normal
    # prose paragraph so we cannot steal content from the following page flow.
    while cursor in records and cursor <= max(indexes) + 3:
        record = records[cursor]
        if int(record.get("omml_count", 0)) > 0:
            if record["id"] not in ids:
                ids.append(record["id"])
            cursor += 1
            continue
        if normalize_text(record.get("text", "")):
            break
        cursor += 1
    return ids


def _set_header_footer(section, header_text: str, footer_text: str) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header
    footer = section.footer
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    _clear_paragraph(hp)
    _clear_paragraph(fp)
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if header_text:
        run = hp.add_run(_clean_pdf_text(header_text))
        run.font.name = "Arial"
        run.font.size = Pt(7.4)
        _set_run_language(run)
    if footer_text:
        run = fp.add_run(_clean_pdf_text(footer_text))
        run.font.name = "Arial"
        run.font.size = Pt(7.2)
        _set_run_language(run)



def _set_section_columns(section, columns: list[dict[str, float]]) -> dict[str, Any]:
    sect_pr = section._sectPr
    existing = sect_pr.find(qn("w:cols"))
    if existing is not None:
        sect_pr.remove(existing)
    cols = OxmlElement("w:cols")
    if len(columns) == 2:
        gutter = max(8.0, float(columns[1]["x0"]) - float(columns[0]["x1"]))
        cols.set(qn("w:num"), "2")
        cols.set(qn("w:space"), str(round(gutter * TWIPS_PER_PT)))
        cols.set(qn("w:equalWidth"), "1")
        result = {
            "policy": "word-section-columns",
            "columnCount": 2,
            "equalWidth": True,
            "gutterPt": round(gutter, 3),
        }
    else:
        cols.set(qn("w:num"), "1")
        result = {
            "policy": "word-section-columns",
            "columnCount": 1,
            "equalWidth": True,
            "gutterPt": None,
        }
    sect_pr.append(cols)
    return result


def _add_spanning_text_frame(
    doc: Document,
    item: dict[str, Any],
    matches: dict[str, dict[str, Any]],
    docx_paras: dict[str, dict[str, Any]],
    body_size: float,
) -> None:
    text, source_runs, _, _ = _item_text_and_runs(item, matches, docx_paras)
    if not text.strip():
        return
    p = doc.add_paragraph()
    _set_frame(p, list(map(float, item["bbox"])))
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Pt(0)
    p.paragraph_format.right_indent = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(body_size * 1.18)
    sem = item.get("semantic_type", "body")
    if sem == "heading":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_runs(p, text, source_runs, max(body_size, 10.5), color=RGBColor(0,0,255), force_bold=True, preserve_line_breaks=True)
    elif sem == "caption":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_runs(p, text, source_runs, max(5.0, body_size * 0.9), color=RGBColor(55,95,135), italic=True)
    else:
        _add_runs(p, text, source_runs, body_size)

def _configure_section(section, page_struct: dict[str, Any]) -> dict[str, Any]:
    width = float(page_struct["width_pt"])
    height = float(page_struct["height_pt"])
    col = page_struct["main_column"]
    footer_y = min((float(f["bbox"][1]) for f in page_struct.get("footers", [])), default=height - 112.0)
    explicit_margins = page_struct.get("section_margins") or {}
    left_margin = float(explicit_margins.get("left") or max(18.0, float(col["x0"])))
    right_margin = float(explicit_margins.get("right") or max(18.0, width - float(col["x1"])))
    top_margin = float(explicit_margins.get("top") or max(36.0, float(col["y0"])))
    bottom_margin = float(explicit_margins.get("bottom") or max(34.0, height - footer_y + 2.0))
    section.page_width = Pt(width)
    section.page_height = Pt(height)
    section.left_margin = Pt(left_margin)
    section.right_margin = Pt(right_margin)
    section.top_margin = Pt(top_margin)
    section.bottom_margin = Pt(bottom_margin)
    section.header_distance = Pt(15)
    section.footer_distance = Pt(14)
    column_config = _set_section_columns(section, list(page_struct.get("columns", [])))
    return {
        "pageWidthPt": round(width, 3),
        "pageHeightPt": round(height, 3),
        "leftMarginPt": round(left_margin, 3),
        "rightMarginPt": round(right_margin, 3),
        "topMarginPt": round(top_margin, 3),
        "bottomMarginPt": round(bottom_margin, 3),
        "marginSource": page_struct.get("section_margin_source") or ("explicit-section-margins" if explicit_margins else "pdf-main-column"),
        "mirrorMargins": bool(explicit_margins.get("mirror")),
        **column_config,
    }


def _scaled_docx_section_margins(
    docx_analysis: dict[str, Any],
    *,
    target_width: float,
    target_height: float,
) -> dict[str, Any] | None:
    sections = [
        section for section in docx_analysis.get("sections", []) or []
        if float(section.get("page_width_pt") or 0.0) > 0 and float(section.get("page_height_pt") or 0.0) > 0
    ]
    if not sections:
        return None
    section = sections[0]
    source_width = float(section.get("page_width_pt") or target_width)
    source_height = float(section.get("page_height_pt") or target_height)
    sx = target_width / source_width if source_width else 1.0
    sy = target_height / source_height if source_height else 1.0
    margins = {
        "left": max(18.0, float(section.get("left_margin_pt") or 72.0) * sx),
        "right": max(18.0, float(section.get("right_margin_pt") or 72.0) * sx),
        "top": max(36.0, float(section.get("top_margin_pt") or 72.0) * sy),
        "bottom": max(34.0, float(section.get("bottom_margin_pt") or 72.0) * sy),
    }
    if margins["left"] + margins["right"] >= target_width * 0.78:
        return None
    if margins["top"] + margins["bottom"] >= target_height * 0.55:
        return None
    return {
        "source": "docx-section-scaled-to-pdf-page",
        "sourceSection": section,
        "scaleX": round(sx, 5),
        "scaleY": round(sy, 5),
        "margins": margins,
    }


def _flow_section_page_struct(
    pages: list[dict[str, Any]],
    docx_analysis: dict[str, Any],
    page_layout_spine: dict[str, Any] | None = None,
) -> dict[str, Any]:
    if not pages:
        raise ValueError("Cannot derive Word free-flow section without PDF page geometry.")

    widths = [float(page.get("width_pt") or 0.0) for page in pages if page.get("width_pt")]
    heights = [float(page.get("height_pt") or 0.0) for page in pages if page.get("height_pt")]
    width = median(widths) if widths else 595.276
    height = median(heights) if heights else 841.89
    preflight = (page_layout_spine or {}).get("layoutPreflight") or {}
    preflight_page_setup = preflight.get("pageSetupEstimate") or {}
    if not preflight_page_setup:
        raise ValueError("Missing page_layout_spine.layoutPreflight.pageSetupEstimate; maps-first build cannot continue.")

    docx_margin_record = _scaled_docx_section_margins(docx_analysis, target_width=width, target_height=height)
    width = float(preflight_page_setup.get("pageWidthPt") or width)
    height = float(preflight_page_setup.get("pageHeightPt") or height)
    left = float(preflight_page_setup.get("leftMarginPt") or 72.0)
    right = float(preflight_page_setup.get("rightMarginPt") or 72.0)
    top = float(preflight_page_setup.get("topMarginPt") or 72.0)
    bottom = float(preflight_page_setup.get("bottomMarginPt") or 72.0)
    margin_source = str(preflight_page_setup.get("marginSource") or "page-layout-spine-preflight")
    margin_evidence = {
        "source": "page_layout_spine.layoutPreflight",
        "pageSetupEstimate": preflight_page_setup,
        "columnProfile": preflight.get("columnProfile"),
        "localTypographyPolicy": preflight.get("localTypographyPolicy"),
        "docxSectionCandidate": docx_margin_record,
    }
    mirror = bool(preflight_page_setup.get("mirrorMargins"))
    return {
        "page": "document-flow",
        "width_pt": width,
        "height_pt": height,
        "section_margins": {
            "left": left,
            "right": right,
            "top": top,
            "bottom": bottom,
            "mirror": mirror,
        },
        "section_margin_source": margin_source,
        "section_margin_evidence": margin_evidence,
        "main_column": {
            "x0": left,
            "x1": max(left + 120.0, width - right),
            "y0": top,
            "y1": max(top + 120.0, height - bottom),
        },
        "columns": [],
        "headers": [],
        "footers": [],
        "layout_mode": "single_column",
    }


def _docx_para_map(docx_analysis: dict[str, Any]) -> dict[str, dict[str, Any]]:
    return {p["id"]: p for p in docx_analysis.get("paragraphs", [])}


def _match_map(alignment: dict[str, Any]) -> dict[str, dict[str, Any]]:
    return {m["pdf_region"]: m for m in alignment.get("matches", [])}


def _item_text_and_runs(
    item: dict[str, Any],
    matches: dict[str, dict[str, Any]],
    docx_paras: dict[str, dict[str, Any]],
) -> tuple[str, list[dict[str, Any]] | None, str, list[str]]:
    texts: list[str] = []
    all_runs: list[dict[str, Any]] = []
    paragraph_ids: list[str] = []
    all_strong = True
    source = "pdf"
    for rid in item.get("region_ids", []):
        match = matches.get(rid)
        if not match or match.get("status") != "strong" or not match.get("docx_text"):
            all_strong = False
            break
        texts.append(str(match["docx_text"]).strip())
        for pos, pid in enumerate(match.get("docx_paragraphs", [])):
            if pid not in paragraph_ids:
                paragraph_ids.append(pid)
            record = docx_paras.get(pid)
            if record:
                all_runs.extend(record.get("runs", []))
                if pos < len(match.get("docx_paragraphs", [])) - 1:
                    all_runs.append({"text": "\n", "bold": None, "italic": None, "underline": None})
    if all_strong and texts:
        source = "docx-strong"
        return "\n".join(texts), all_runs or None, source, paragraph_ids
    return _clean_pdf_text(item.get("text", "")), None, source, []


def _add_runs(
    paragraph,
    text: str,
    source_runs: list[dict[str, Any]] | None,
    font_size: float,
    color: RGBColor | None = None,
    force_bold: bool = False,
    italic: bool = False,
    preserve_line_breaks: bool = False,
) -> None:
    if source_runs:
        for src in source_runs:
            value = str(src.get("text", ""))
            if not value:
                continue
            if not preserve_line_breaks:
                value = re.sub(r"\s+", " ", value)
            run = paragraph.add_run(value)
            run.bold = force_bold or bool(src.get("bold"))
            run.italic = italic or bool(src.get("italic"))
            if src.get("underline"):
                run.underline = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(font_size)
            if color:
                run.font.color.rgb = color
            _set_run_language(run)
    else:
        parts = text.split("\n") if preserve_line_breaks else [re.sub(r"\s+", " ", text).strip()]
        for i, part in enumerate(parts):
            run = paragraph.add_run(part)
            run.bold = force_bold
            run.italic = italic
            run.font.name = "Times New Roman"
            run.font.size = Pt(font_size)
            if color:
                run.font.color.rgb = color
            _set_run_language(run)
            if preserve_line_breaks and i < len(parts) - 1:
                run.add_break(WD_BREAK.LINE)


def _regions_by_page(pdf_analysis: dict[str, Any]) -> dict[int, dict[str, dict[str, Any]]]:
    result: dict[int, dict[str, dict[str, Any]]] = {}
    for page in pdf_analysis.get("pages", []):
        result[int(page["page"])] = {r["id"]: r for r in page.get("regions", [])}
    return result


def _layout_spine_order(page_layout_spine: dict[str, Any] | None, item: dict[str, Any], page_no: int | None = None) -> int | None:
    if not page_layout_spine:
        return None
    order_map = page_layout_spine.get("layoutOrderBySlot") or {}
    has_page_qualified_keys = any(":" in str(key) for key in order_map.keys())
    keys = [item.get("id"), item.get("visual_group_id")]
    if item.get("id") and str(item.get("id")).startswith("flow-"):
        keys.append(str(item.get("id"))[5:])
    for key in keys:
        if key is None:
            continue
        value = None
        if page_no is not None:
            value = order_map.get(f"{page_no}:{key}")
        if value is None and not has_page_qualified_keys:
            value = order_map.get(str(key))
        if value is not None:
            try:
                return int(value)
            except (TypeError, ValueError):
                return None
    return None


def _sort_flow_with_layout_spine(
    flow_items: list[dict[str, Any]],
    page_layout_spine: dict[str, Any] | None,
    page_no: int | None = None,
) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    audit = {
        "page": page_no,
        "policy": "pdf-geometry-order",
        "reason": "no-layout-spine",
        "item_count": len(flow_items),
        "ordered_item_count": 0,
        "ordered_ratio": 0.0,
        "backtrack_count": 0,
    }
    if not page_layout_spine:
        return flow_items, audit
    ordered_count = sum(1 for item in flow_items if _layout_spine_order(page_layout_spine, item, page_no) is not None)
    ordered_ratio = ordered_count / max(1, len(flow_items))
    def key(item: dict[str, Any]) -> tuple[int, int, float, float]:
        order = _layout_spine_order(page_layout_spine, item, page_no)
        box = item.get("bbox") or [0, 0, 0, 0]
        try:
            y0 = float(box[1])
            x0 = float(box[0])
        except Exception:
            y0 = 0.0
            x0 = 0.0
        return (0 if order is not None else 1, int(order if order is not None else 10**9), y0, x0)
    candidate = sorted(flow_items, key=key)
    backtracks = 0
    max_backtrack = 0.0
    last_y0_by_column: dict[int, float] = {}
    two_columns = len({int(item.get("column_index", 0) or 0) for item in flow_items}) > 1
    for item in candidate:
        column = int(item.get("column_index", 0) or 0) if two_columns else 0
        box = item.get("bbox") or [0, 0, 0, 0]
        try:
            y0 = float(box[1])
        except Exception:
            y0 = 0.0
        previous_y0 = last_y0_by_column.get(column)
        if previous_y0 is not None and y0 + 4.0 < previous_y0:
            backtracks += 1
            max_backtrack = max(max_backtrack, previous_y0 - y0)
        last_y0_by_column[column] = max(previous_y0 if previous_y0 is not None else y0, y0)
    audit.update({
        "ordered_item_count": ordered_count,
        "ordered_ratio": round(ordered_ratio, 4),
        "backtrack_count": backtracks,
        "max_backtrack_pt": round(max_backtrack, 2),
    })
    if ordered_ratio >= 0.85 and backtracks == 0:
        audit.update({"policy": "layout-spine-order", "reason": "page-layout-spine-monotonic"})
        return candidate, audit
    audit["reason"] = "layout-spine-order-rejected"
    return flow_items, audit


def _word_flow_gap(raw_gap: float, item: dict[str, Any], body_size: float, gap_scale: float) -> dict[str, Any]:
    scaled = max(0.0, raw_gap) * max(0.0, float(gap_scale))
    semantic_type = str(item.get("semantic_type") or item.get("type") or "body")
    if semantic_type == "heading":
        limit = max(4.0, body_size * 0.85)
    elif semantic_type in {"figure", "equation"} or item.get("type") == "visual":
        limit = max(3.0, body_size * 0.55)
    elif semantic_type == "caption":
        limit = max(2.0, body_size * 0.45)
    else:
        limit = max(3.0, body_size * 0.65)
    clamped = min(scaled, limit)
    natural_stops = [0.0, 1.5, 3.0, 6.0, 8.0, 12.0]
    candidates = [value for value in natural_stops if value <= limit + 0.01]
    if not candidates:
        candidates = [0.0]
    applied = min(candidates, key=lambda value: (abs(value - clamped), value))
    return {
        "raw_gap_pt": round(max(0.0, raw_gap), 2),
        "scaled_gap_pt": round(scaled, 2),
        "applied_gap_pt": round(applied, 2),
        "gap_clamped": clamped + 0.01 < scaled,
        "gap_quantized": abs(applied - clamped) > 0.01,
        "gap_policy": "word-flow-natural-gap",
    }


def build_native_page_document(
    pdf_analysis: dict[str, Any],
    page_structure: dict[str, Any],
    alignment: dict[str, Any],
    docx_analysis: dict[str, Any],
    style_profile: dict[str, Any],
    output_path: Path,
    body_size_override: float | None = None,
    font_scale: float = 1.0,
    gap_scale: float = 0.72,
    body_line_spacing_multiple: float | None = None,
    docx_donor_map: dict[str, Any] | None = None,
    page_layout_spine: dict[str, Any] | None = None,
    flow_mode: str = "free",
) -> dict[str, Any]:
    """Build a native DOCX page-structure probe for an arbitrary page range.

    Main prose is normal Word flow. Only actual side callouts become positioned
    paragraph frames. Figures and equations are grouped into one crop per logical
    visual object and inserted inline or as floating anchored pictures. No layout
    tables and no full-page image backgrounds are used.
    """
    doc = Document()
    _set_cell_free_document_settings(doc)
    font_scale = max(0.5, min(1.5, float(font_scale or 1.0)))
    raw_body_size = float(body_size_override or style_profile.get("inferred_body_font_size_pt") or 10.5)
    body_size = raw_body_size * font_scale
    body_line_spacing = (
        float(body_line_spacing_multiple)
        if body_line_spacing_multiple is not None and float(body_line_spacing_multiple) > 0
        else None
    )
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(body_size)
    normal.paragraph_format.space_after = Pt(0)
    if body_line_spacing is not None:
        normal.paragraph_format.line_spacing = body_line_spacing
    else:
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        normal.paragraph_format.line_spacing = Pt(body_size * 1.25)

    matches = _match_map(alignment)
    docx_paras = _docx_para_map(docx_analysis)
    source_doc, source_paragraphs = _source_paragraph_map(Path(docx_analysis["source"]))
    regions_by_page = _regions_by_page(pdf_analysis)
    build_report: dict[str, Any] = {
        "body_size_pt": round(body_size, 3),
        "font_scale": round(font_scale, 4),
        "gap_scale": round(float(gap_scale), 3),
        "body_line_spacing_multiple": round(body_line_spacing, 3) if body_line_spacing is not None else None,
        "body_line_spacing_policy": "word-natural-multiple" if body_line_spacing is not None else "pdf-line-pitch",
        "layout_spine": {
            "version": (page_layout_spine or {}).get("version"),
            "coverage": ((page_layout_spine or {}).get("summary") or {}).get("coverage"),
            "contract_coverage": ((page_layout_spine or {}).get("summary") or {}).get("contractCoverage"),
            "safe_flow_ordering_slots": ((page_layout_spine or {}).get("summary") or {}).get("safeFlowOrderingSlotCount"),
            "policy": "guarded-layout-spine-order-for-word-flow" if page_layout_spine else "page-structure-order",
        },
        "flow_order_audit": [],
        "flow_geometry_fit_summary": {
            "policy": "pdf-paragraph-geometry-fit-before-render",
            "count": 0,
        },
        "pages": [],
    }
    flow_geometry_fits: list[dict[str, Any]] = []
    used_native_math_indexes: set[int] = set()

    pages = page_structure.get("pages", [])
    use_free_flow = str(flow_mode or "free").strip().lower() in {"free", "word-free-flow", "single-section"}
    if not use_free_flow:
        raise ValueError("Only word-free-flow single-section reconstruction is supported.")
    flow_section_page_struct = _flow_section_page_struct(list(pages), docx_analysis, page_layout_spine)
    _set_mirror_margins(doc, bool((flow_section_page_struct.get("section_margins") or {}).get("mirror")))
    flow_section_config = _configure_section(doc.sections[-1], flow_section_page_struct)
    build_report["section_break_policy"] = "word-free-flow-single-section"
    build_report["word_section_policy"] = {
        "policy": build_report["section_break_policy"],
        "hardBreaksBetweenPdfPages": False,
        "sectionPerPdfPage": False,
        "intendedSectionCount": 1,
        "marginsSource": flow_section_page_struct.get("section_margin_source"),
        "paginationAuthority": "microsoft-word-natural-flow",
        "layoutPreflight": {
            "source": "pdf-page-structure-before-render",
            "pageCount": len(pages),
            "twoColumnPageCount": sum(1 for page in pages if page.get("layout_mode") == "two_columns"),
            "singleColumnPageCount": sum(1 for page in pages if page.get("layout_mode") != "two_columns"),
            "localFontPolicy": "pdf-span-dominant-size-per-flow-item",
            "lineHeightPolicy": "pdf-line-pitch-per-flow-item",
        },
        "section_config": flow_section_config,
        "section_margin_evidence": flow_section_page_struct.get("section_margin_evidence"),
    }
    build_report["layout_spine"]["policy"] = (
        "word-free-flow-single-section"
        if page_layout_spine
        else build_report["layout_spine"]["policy"]
    )
    for page_index, page_struct in enumerate(pages):
        section_config = flow_section_config or {}

        all_flow_items = page_struct.get("flow", [])
        _assign_native_math_sources(
            all_flow_items, matches, docx_analysis,
            alignment.get("summary", {}).get("candidate_docx_paragraph_range"),
            used_native_math_indexes,
            docx_donor_map,
        )
        spanning_items = [item for item in all_flow_items if item.get("spanning")]
        flow_items = [item for item in all_flow_items if not item.get("spanning")]
        flow_items, flow_order_audit = _sort_flow_with_layout_spine(flow_items, page_layout_spine, int(page_struct["page"]))
        build_report["flow_order_audit"].append(flow_order_audit)

        # A caption that belongs to a page-positioned figure must not participate
        # in normal flow.  Otherwise Word may push the caption to a new page even
        # though the source page has a reserved caption rectangle directly below
        # the floating figure.  Keep it editable as a page-relative text frame.
        floating_groups = [
            group for group in page_struct.get("visual_groups", [])
            if group.get("placement") == "floating"
        ]
        positioned_captions: list[dict[str, Any]] = []
        remaining_flow_items: list[dict[str, Any]] = []
        for item in flow_items:
            if item.get("type") != "text" or item.get("semantic_type") != "caption":
                remaining_flow_items.append(item)
                continue
            ix0, iy0, ix1, iy1 = map(float, item.get("bbox", [0, 0, 0, 0]))
            associated = False
            for group in floating_groups:
                gx0, gy0, gx1, gy1 = map(float, group.get("bbox", [0, 0, 0, 0]))
                horizontal_overlap = max(0.0, min(ix1, gx1) - max(ix0, gx0))
                caption_width = max(1.0, ix1 - ix0)
                # Captions may overlap the bottom few points of the figure in the
                # PDF.  A small downward separation is also accepted.
                if horizontal_overlap / caption_width >= 0.45 and -12.0 <= iy0 - gy1 <= 24.0:
                    associated = True
                    break
            if associated:
                positioned_captions.append(item)
            else:
                remaining_flow_items.append(item)
        flow_items = remaining_flow_items

        # Full-width headings/instructions above a two-column exercise block and
        # captions tied to floating figures are positioned as native text frames.
        # They remain editable and do not consume main-flow page height.
        for item in [*spanning_items, *positioned_captions]:
            if item.get("type") == "text":
                _add_spanning_text_frame(doc, item, matches, docx_paras, body_size)

        # Floating page objects first. Their anchor paragraphs have one-point line height.
        z = 251658240
        for group in page_struct.get("visual_groups", []):
            if group.get("placement") != "floating":
                continue
            crop_path = Path(group["crop_path"])
            svg_path = Path(group["svg_path"]) if group.get("svg_path") else None
            _add_floating_picture(doc, crop_path, group["bbox"], wrap=group.get("wrap", "none"), z=z, svg_path=svg_path)
            z += 1

        page_regions = regions_by_page[int(page_struct["page"])]
        callout_builds: list[dict[str, Any]] = []
        for callout in page_struct.get("callouts", []):
            match = matches.get(callout["id"])
            text = str(match.get("docx_text")) if match and match.get("status") == "strong" else _clean_pdf_text(callout.get("text", ""))
            paragraph_ids = _callout_source_paragraph_ids(callout, match, docx_analysis)
            callout_builds.append(_add_callout(
                doc, callout, text, body_size,
                source_paragraphs=source_paragraphs if paragraph_ids else None,
                paragraph_ids=paragraph_ids,
            ))

        previous_y1_by_column: dict[int, float | None] = {0: None, 1: None}
        current_column = 0
        two_columns = False
        column_break_count = 0
        page_items: list[dict[str, Any]] = []
        for item in flow_items:
            item_column = int(item.get("column_index", 0) or 0) if two_columns else 0
            if two_columns and item_column != current_column:
                breaker = doc.add_paragraph()
                breaker.paragraph_format.space_before = Pt(0)
                breaker.paragraph_format.space_after = Pt(0)
                breaker.add_run().add_break(WD_BREAK.COLUMN)
                column_break_count += 1
                current_column = item_column
            y0, y1 = float(item["bbox"][1]), float(item["bbox"][3])
            previous_y1 = previous_y1_by_column.get(item_column)
            raw_gap = 0.0 if previous_y1 is None else max(0.0, y0 - previous_y1)
            gap_info = _word_flow_gap(raw_gap, item, body_size, gap_scale)
            gap = float(gap_info["applied_gap_pt"])
            if item["type"] == "visual":
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(gap)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.keep_together = True
                native_math_ids = list(item.get("native_math_paragraph_ids", []))
                if item.get("semantic_type") == "equation" and native_math_ids:
                    native_info = _append_source_paragraph_content(
                        p, source_paragraphs, native_math_ids, body_size,
                        flatten_math_paragraphs=True,
                    )
                    page_items.append({
                        "id": item["id"],
                        "type": item["semantic_type"],
                        "source": "docx-native-omml",
                        "native_math_layout_policy": "flatten-display-omml-in-visual-flow-slot",
                        "docx_paragraphs": native_math_ids,
                        "math_count": native_info["math_count"],
                        "native_math_match_score": item.get("native_math_match_score"),
                        "native_math_text_score": item.get("native_math_text_score"),
                        "native_math_signature_source": item.get("native_math_signature_source"),
                        "gap_pt": round(gap, 2),
                        **gap_info,
                    })
                elif item.get("semantic_type") == "equation" and item.get("markdown_latex"):
                    recovered = _append_latex_as_omml(p, item.get("markdown_latex"), item)
                    if recovered:
                        page_items.append({
                            "id": item["id"],
                            "type": item["semantic_type"],
                            "source": "markdown-latex-to-omml",
                            "math_count": recovered,
                            "donor": item.get("markdown_equation_donor"),
                            "gap_pt": round(gap, 2),
                            **gap_info,
                        })
                    else:
                        run = p.add_run()
                        box = item["bbox"]
                        _add_picture_with_optional_svg(run, Path(item["crop_path"]), None, width=Pt(float(box[2])-float(box[0])), height=Pt(float(box[3])-float(box[1])))
                        page_items.append({
                            "id": item["id"],
                            "type": item["semantic_type"],
                            "source": "page-crop-after-latex-conversion-failure",
                            "latex_omml_failure": item.get("latex_omml_failure"),
                            "gap_pt": round(gap, 2),
                            **gap_info,
                        })
                else:
                    box = item["bbox"]
                    width = float(box[2]) - float(box[0])
                    height = float(box[3]) - float(box[1])
                    if two_columns:
                        active_col = page_struct["columns"][item_column]
                        col_width = float(active_col["x1"]) - float(active_col["x0"])
                    else:
                        col_width = float(page_struct["main_column"]["x1"]) - float(page_struct["main_column"]["x0"])
                    scale = min(1.0, col_width / max(1.0, width))
                    run = p.add_run()
                    svg_path = Path(item["svg_path"]) if item.get("svg_path") else None
                    _add_picture_with_optional_svg(
                        run, Path(item["crop_path"]), svg_path,
                        width=Pt(width * scale), height=Pt(height * scale),
                    )
                    page_items.append({
                        "id": item["id"],
                        "type": item["semantic_type"],
                        "source": "mathpix-native-svg" if svg_path else str(item.get("asset_source") or "page-crop"),
                        "gap_pt": round(gap, 2),
                        **gap_info,
                    })
            else:
                semantic_type = item.get("semantic_type", "body")
                text, source_runs, source, paragraph_ids = _item_text_and_runs(item, matches, docx_paras)
                has_native_math = any(int(docx_paras.get(pid, {}).get("omml_count", 0)) > 0 for pid in paragraph_ids)
                if not text.strip() and not has_native_math:
                    continue
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(gap)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.keep_together = semantic_type in {"heading", "caption"}
                p.paragraph_format.widow_control = True
                region_ids = item.get("region_ids", [])
                max_size = _max_span_size(page_regions, region_ids, raw_body_size) * font_scale
                native_info = None
                line_height = None
                line_height_policy = "pdf-line-pitch-half-point"
                if semantic_type == "heading":
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    size = max(body_size, min(12.5, max_size))
                    line_height = max(size * 1.15, _line_height(page_regions, region_ids, size))
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    p.paragraph_format.line_spacing = Pt(line_height)
                    if normalize_text(text).startswith(("εφαρμογη", "εφαρμογή")):
                        _set_paragraph_borders(p, color="808080", fill="E7E6E6", size=6)
                        p.paragraph_format.left_indent = Pt(4)
                        p.paragraph_format.right_indent = Pt(4)
                        p.paragraph_format.space_before = Pt(max(0.0, gap))
                    if has_native_math:
                        native_info = _append_source_paragraph_content(
                            p, source_paragraphs, paragraph_ids, size,
                            color=RGBColor(0, 0, 255), force_bold=True,
                        )
                    else:
                        _add_runs(p, text, source_runs, size, color=RGBColor(0, 0, 255), force_bold=True, preserve_line_breaks=True)
                elif semantic_type == "caption":
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    size = max(5.0, min(body_size * 0.92, max_size))
                    line_height = max(size * 1.15, _line_height(page_regions, region_ids, size))
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    p.paragraph_format.line_spacing = Pt(line_height)
                    if has_native_math:
                        native_info = _append_source_paragraph_content(
                            p, source_paragraphs, paragraph_ids, size,
                            color=RGBColor(55, 95, 135), italic=True,
                        )
                    else:
                        _add_runs(p, text, source_runs, size, color=RGBColor(55, 95, 135), italic=True)
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    size = max(5.0, min(12.5, _dominant_span_size(page_regions, region_ids, raw_body_size) * font_scale))
                    if body_line_spacing is not None:
                        line_height = round((size * body_line_spacing) * 2.0) / 2.0
                        p.paragraph_format.line_spacing = body_line_spacing
                        line_height_policy = "word-natural-multiple"
                    else:
                        line_height = _line_height(page_regions, region_ids, size)
                        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                        p.paragraph_format.line_spacing = Pt(line_height)
                        line_height_policy = "pdf-line-pitch-half-point"
                    if has_native_math:
                        native_info = _append_source_paragraph_content(
                            p, source_paragraphs, paragraph_ids, size
                        )
                        source = "docx-native-omml"
                    else:
                        _add_runs(p, text, source_runs, size)
                if two_columns:
                    active_col = page_struct["columns"][item_column]
                    column_width = float(active_col["x1"]) - float(active_col["x0"])
                else:
                    column_width = float(page_struct["main_column"]["x1"]) - float(page_struct["main_column"]["x0"])
                geometry_fit = _flow_geometry_fit(
                    page_regions,
                    item,
                    column_width_pt=column_width,
                    font_size_pt=size,
                    line_height_pt=float(line_height or 0.0),
                    gap_pt=gap,
                )
                if geometry_fit:
                    flow_geometry_fits.append(geometry_fit)
                page_items.append({
                    "id": item["id"],
                    "type": semantic_type,
                    "source": source,
                    "docx_paragraphs": paragraph_ids,
                    "native_math_count": int(native_info["math_count"]) if native_info else 0,
                    "font_size_pt": round(size, 2),
                    "line_height_pt": round(float(line_height or 0.0), 2),
                    "line_height_policy": line_height_policy,
                    "line_spacing_multiple": (
                        round(body_line_spacing, 3)
                        if semantic_type not in {"heading", "caption"} and body_line_spacing is not None
                        else None
                    ),
                    "gap_pt": round(gap, 2),
                    "flow_geometry_fit": geometry_fit,
                    **gap_info,
                })
            previous_y1_by_column[item_column] = max(previous_y1 or y1, y1)

        build_report["pages"].append({
            "page": page_struct["page"],
            "section_config": section_config,
            "main_column": page_struct["main_column"],
            "flow_item_count": len(page_struct.get("flow", [])),
            "layout_mode": page_struct.get("layout_mode", "single_column"),
            "column_count": len(page_struct.get("columns", [])) or 1,
            "column_flow_contract": {
                "policy": "disabled-in-document-free-flow",
                "active": bool(two_columns),
                "left_item_count": sum(1 for item in flow_items if item.get("column_index") == 0 and not item.get("spanning")),
                "right_item_count": sum(1 for item in flow_items if item.get("column_index") == 1 and not item.get("spanning")),
                "column_break_count": column_break_count,
                "passes": (not two_columns) or column_break_count == 1,
            },
            "spanning_item_count": len(spanning_items),
            "positioned_caption_count": len(positioned_captions),
            "callout_count": len(page_struct.get("callouts", [])),
            "callout_builds": callout_builds,
            "floating_visual_count": sum(1 for g in page_struct.get("visual_groups", []) if g.get("placement") == "floating"),
            "inline_visual_count": sum(1 for g in page_struct.get("visual_groups", []) if g.get("placement") == "inline"),
            "items": page_items,
        })

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    build_report["flow_geometry_fit_summary"] = _summarize_flow_geometry_fits(flow_geometry_fits)
    return build_report
