from __future__ import annotations

import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from statistics import median
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


VERSION = "native-builder-canonical-0.1.2"
TWIPS_PER_PT = 20
PAGE_WIDTH_PT = 595.276
PAGE_HEIGHT_PT = 841.89


def _rgb_hex(value: Any) -> str | None:
    if isinstance(value, str):
        text = value.strip().lstrip("#")
        if len(text) == 6 and all(ch in "0123456789abcdefABCDEF" for ch in text):
            return text.upper()
    try:
        number = int(value)
    except (TypeError, ValueError):
        return None
    return f"{number & 0xFFFFFF:06X}"


def _page_scale(canonical: dict[str, Any]) -> tuple[float, float, float, float]:
    page_evidence = canonical.get("pageEvidence") or {}
    size = page_evidence.get("pageSizePx") or page_evidence.get("pageDimensionsPx") or {}

    width_px = 0.0
    height_px = 0.0
    if isinstance(size, dict):
        width_px = float(size.get("width") or size.get("widthPx") or 0.0)
        height_px = float(size.get("height") or size.get("heightPx") or 0.0)
    elif isinstance(size, (list, tuple)) and len(size) >= 2:
        width_px = float(size[0] or 0.0)
        height_px = float(size[1] or 0.0)
    elif size:
        raise RuntimeError(
            f"Canonical Word build blocked: unsupported pageSizePx schema: {type(size).__name__}"
        )

    if width_px <= 0 or height_px <= 0:
        blocks = canonical.get("blocks") or []
        xs = [
            float((b.get("geometry") or {}).get("bboxPx", [0, 0, 0, 0])[2])
            for b in blocks
            if isinstance((b.get("geometry") or {}).get("bboxPx"), (list, tuple))
            and len((b.get("geometry") or {}).get("bboxPx")) == 4
        ]
        ys = [
            float((b.get("geometry") or {}).get("bboxPx", [0, 0, 0, 0])[3])
            for b in blocks
            if isinstance((b.get("geometry") or {}).get("bboxPx"), (list, tuple))
            and len((b.get("geometry") or {}).get("bboxPx")) == 4
        ]
        width_px = max(max(xs, default=0.0), 2067.0)
        height_px = max(max(ys, default=0.0), 2924.0)

    return width_px, height_px, PAGE_WIDTH_PT / width_px, PAGE_HEIGHT_PT / height_px


def _bbox_pt(box: Any, sx: float, sy: float) -> list[float] | None:
    if not isinstance(box, (list, tuple)) or len(box) != 4:
        return None
    try:
        values = [
            float(box[0]) * sx,
            float(box[1]) * sy,
            float(box[2]) * sx,
            float(box[3]) * sy,
        ]
    except (TypeError, ValueError):
        return None
    if values[2] <= values[0] or values[3] <= values[1]:
        return None
    return [round(v, 3) for v in values]


def _set_language(run, lang: str = "el-GR") -> None:
    r_pr = run._r.get_or_add_rPr()
    node = r_pr.find(qn("w:lang"))
    if node is None:
        node = OxmlElement("w:lang")
        r_pr.append(node)
    node.set(qn("w:val"), lang)
    node.set(qn("w:eastAsia"), lang)


def _set_unequal_columns(section, columns: list[dict[str, Any]]) -> dict[str, Any]:
    sect_pr = section._sectPr
    existing = sect_pr.find(qn("w:cols"))
    if existing is not None:
        sect_pr.remove(existing)
    cols = OxmlElement("w:cols")
    if len(columns) != 2:
        cols.set(qn("w:num"), "1")
        sect_pr.append(cols)
        return {"columnCount": 1, "equalWidth": True, "columns": []}

    left, right = columns
    left_w = max(1.0, float(left["x1"]) - float(left["x0"]))
    right_w = max(1.0, float(right["x1"]) - float(right["x0"]))
    gutter = max(0.0, float(right["x0"]) - float(left["x1"]))
    cols.set(qn("w:num"), "2")
    cols.set(qn("w:equalWidth"), "0")
    cols.set(qn("w:space"), str(round(gutter * TWIPS_PER_PT)))

    c1 = OxmlElement("w:col")
    c1.set(qn("w:w"), str(round(left_w * TWIPS_PER_PT)))
    c1.set(qn("w:space"), str(round(gutter * TWIPS_PER_PT)))
    c2 = OxmlElement("w:col")
    c2.set(qn("w:w"), str(round(right_w * TWIPS_PER_PT)))
    c2.set(qn("w:space"), "0")
    cols.append(c1)
    cols.append(c2)
    sect_pr.append(cols)

    return {
        "columnCount": 2,
        "equalWidth": False,
        "gutterPt": round(gutter, 3),
        "columns": [
            {"widthPt": round(left_w, 3), "x0": left["x0"], "x1": left["x1"]},
            {"widthPt": round(right_w, 3), "x0": right["x0"], "x1": right["x1"]},
        ],
    }


def _paragraph_frame_style(paragraph, container: dict[str, Any]) -> dict[str, Any]:
    p_pr = paragraph._p.get_or_add_pPr()
    stroke = container.get("stroke") if isinstance(container.get("stroke"), dict) else {}
    fill = container.get("fill") if isinstance(container.get("fill"), dict) else {}
    audit = {"borderApplied": False, "fillApplied": False}

    color = _rgb_hex(stroke.get("color"))
    width = stroke.get("widthPt") or stroke.get("width")
    if color and stroke.get("status") != "none":
        try:
            width_pt = max(0.25, float(width or 0.75))
        except (TypeError, ValueError):
            width_pt = 0.75
        borders = OxmlElement("w:pBdr")
        for edge in ("top", "left", "bottom", "right"):
            node = OxmlElement(f"w:{edge}")
            node.set(qn("w:val"), "single")
            node.set(qn("w:sz"), str(max(2, round(width_pt * 8.0))))
            node.set(qn("w:space"), "1")
            node.set(qn("w:color"), color)
            borders.append(node)
        p_pr.append(borders)
        audit.update({
            "borderApplied": True,
            "borderColor": f"#{color}",
            "borderWidthPt": width_pt,
        })

    fill_color = _rgb_hex(fill.get("color"))
    if fill_color and fill.get("status") != "none":
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_color)
        p_pr.append(shd)
        audit.update({"fillApplied": True, "fillColor": f"#{fill_color}"})
    return audit


def _latex_body(text: str) -> str:
    value = str(text or "").strip()
    value = re.sub(r"^\\\[|\\\]$", "", value.strip(), flags=re.S).strip()
    return value


def _append_omml(paragraph, latex: str) -> bool:
    latex = _latex_body(latex)
    if not latex or shutil.which("pandoc") is None:
        return False
    from copy import deepcopy
    import zipfile
    from lxml import etree

    with tempfile.TemporaryDirectory(prefix="canonical-omml-") as td:
        root = Path(td)
        md = root / "eq.md"
        out = root / "eq.docx"
        md.write_text("$$\n" + latex + "\n$$\n", encoding="utf-8")
        try:
            completed = subprocess.run(
                ["pandoc", str(md), "-o", str(out)],
                capture_output=True,
                text=True,
                timeout=30,
            )
        except Exception:
            return False
        if completed.returncode != 0 or not out.exists():
            return False
        with zipfile.ZipFile(out) as archive:
            xml = etree.fromstring(archive.read("word/document.xml"))
        maths = xml.xpath(".//*[local-name()='oMath']")
        if not maths:
            return False
        for math in maths:
            paragraph._p.append(deepcopy(math))
        return True


def _resolve_asset(package_root: Path | None, block: dict[str, Any]) -> Path | None:
    if package_root is None:
        return None
    text = str((block.get("content") or {}).get("text") or "")
    match = re.search(r"([0-9a-fA-F-]{20,}\.(?:png|jpe?g|webp|svg))", text)
    names = [match.group(1)] if match else []
    for candidate in names:
        direct = list(package_root.rglob(candidate))
        if direct:
            return direct[0]

    uuid_match = re.search(r"([0-9a-fA-F]{8}-[0-9a-fA-F-]{27,})", text)
    if uuid_match:
        stem = uuid_match.group(1).lower()
        for path in package_root.rglob("*"):
            if (
                path.is_file()
                and stem in path.name.lower()
                and path.suffix.lower() in {".png", ".jpg", ".jpeg", ".webp", ".svg"}
            ):
                return path
    return None


def _add_word_picture(paragraph, asset: Path, width_pt: float) -> dict[str, Any]:
    """Insert an authoritative canonical asset, normalizing only its image encoding if required."""
    run = paragraph.add_run()
    try:
        run.add_picture(str(asset), width=Pt(width_pt))
        return {"asset": str(asset), "normalization": "not-needed"}
    except Exception as direct_error:
        try:
            from PIL import Image
        except Exception as pil_import_error:
            raise RuntimeError(
                f"Word cannot read canonical image {asset}; Pillow unavailable for encoding normalization: {pil_import_error}"
            ) from direct_error

        try:
            with Image.open(asset) as image:
                detected_format = image.format
                normalized = image.convert("RGB")
                with tempfile.NamedTemporaryFile(
                    prefix="canonical-word-image-",
                    suffix=".png",
                    delete=False,
                ) as handle:
                    normalized_path = Path(handle.name)
                normalized.save(normalized_path, format="PNG")
        except Exception as normalize_error:
            try:
                signature = asset.read_bytes()[:24].hex(" ")
            except Exception:
                signature = "unavailable"
            raise RuntimeError(
                f"Canonical image asset cannot be decoded for Word: {asset}; file signature={signature}; decoder={normalize_error}"
            ) from direct_error

        try:
            paragraph.add_run().add_picture(str(normalized_path), width=Pt(width_pt))
        finally:
            try:
                normalized_path.unlink(missing_ok=True)
            except Exception:
                pass
        return {
            "asset": str(asset),
            "normalization": "pillow-to-png",
            "detectedFormat": detected_format,
        }


def _font_size_pt(block: dict[str, Any], sy: float) -> float:
    values = []
    for value in (block.get("typographyEvidence") or {}).get("fontSizes") or []:
        try:
            number = float(value) * sy
        except (TypeError, ValueError):
            continue
        if number > 0:
            values.append(number)
    return round(median(values), 2) if values else 8.5


def _container_maps(canonical: dict[str, Any]) -> tuple[dict[str, dict[str, Any]], dict[str, str]]:
    containers = {
        str(c.get("id")): c
        for c in canonical.get("pdfContainers") or []
        if c.get("id")
    }
    block_to_container: dict[str, str] = {}
    for container_id, container in containers.items():
        for block_id in container.get("memberBlockIds") or []:
            block_to_container[str(block_id)] = container_id
    return containers, block_to_container


def _zone_order(canonical: dict[str, Any]) -> list[str]:
    topology = canonical.get("pageTopology") or {}
    cross = topology.get("crossZoneReadingOrder") or {}
    order = cross.get("order")
    if isinstance(order, list) and order:
        return [str(v) for v in order]
    raise RuntimeError("Canonical Word build blocked: cross-zone reading order is unresolved.")


def _ordered_blocks(canonical: dict[str, Any], page: int) -> tuple[list[dict[str, Any]], list[str]]:
    topology = canonical.get("pageTopology") or {}
    order = _zone_order(canonical)
    by_id = {str(b.get("id")): b for b in canonical.get("blocks") or []}
    result: list[dict[str, Any]] = []
    for zone_id in order:
        zone = next(
            (z for z in topology.get("zones") or [] if str(z.get("zoneId")) == zone_id),
            None,
        )
        if zone is None:
            raise RuntimeError(f"Canonical Word build blocked: zone missing from topology: {zone_id}")
        for block_id in zone.get("localFlowOrder") or []:
            block = by_id.get(str(block_id))
            if block is None:
                raise RuntimeError(
                    f"Canonical Word build blocked: topology references missing block {block_id}"
                )
            if int((block.get("pageAssignment") or {}).get("physicalPage") or 0) == page:
                result.append(block)
    return result, order


def build_canonical_native_document(
    canonical: dict[str, Any],
    *,
    output_path: Path,
    target_page: int,
    package_root: Path | None = None,
) -> dict[str, Any]:
    topology = canonical.get("pageTopology") or {}
    if int(topology.get("physicalPage") or 0) != int(target_page):
        raise RuntimeError("Canonical Word build blocked: topology page does not match requested page.")
    if (topology.get("crossZoneReadingOrder") or {}).get("status") not in {
        "resolved-by-markdown-record-order",
        "not-needed",
    }:
        raise RuntimeError("Canonical Word build blocked: cross-zone reading order unresolved.")

    width_px, height_px, sx, sy = _page_scale(canonical)
    recovered = topology.get("recoveredFrameEvidence") or {}
    frame = _bbox_pt(recovered.get("bboxPx"), sx, sy)
    if frame is None:
        raise RuntimeError("Canonical Word build blocked: recovered frame unavailable.")

    zone_rows = []
    for zone in topology.get("zones") or []:
        bbox = _bbox_pt(
            zone.get("physicalZoneBBoxPx") or zone.get("canonicalCoverageBBoxPx"),
            sx,
            sy,
        )
        if bbox:
            zone_rows.append({
                "id": str(zone.get("zoneId")),
                "x0": bbox[0],
                "y0": bbox[1],
                "x1": bbox[2],
                "y1": bbox[3],
            })
    if len(zone_rows) != 2:
        raise RuntimeError(
            f"Canonical Word build blocked: page {target_page} requires exactly two resolved zones for this proof, got {len(zone_rows)}."
        )
    zone_rows.sort(key=lambda z: z["x0"])

    ordered, reading_order = _ordered_blocks(canonical, target_page)
    containers, block_to_container = _container_maps(canonical)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Pt(PAGE_WIDTH_PT)
    section.page_height = Pt(PAGE_HEIGHT_PT)
    section.left_margin = Pt(max(0.0, frame[0]))
    section.top_margin = Pt(max(0.0, frame[1]))
    section.right_margin = Pt(max(0.0, PAGE_WIDTH_PT - frame[2]))
    section.bottom_margin = Pt(max(0.0, PAGE_HEIGHT_PT - frame[3]))
    columns_audit = _set_unequal_columns(section, zone_rows)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(8.5)
    normal.paragraph_format.space_after = Pt(0)

    report_items: list[dict[str, Any]] = []
    current_zone = reading_order[0] if reading_order else None
    previous_bottom_by_zone: dict[str, float | None] = {z["id"]: None for z in zone_rows}

    for block in ordered:
        block_id = str(block.get("id"))
        semantic = str((block.get("semantic") or {}).get("type") or "paragraph")
        content = block.get("content") or {}
        text = str(content.get("text") or "")
        geometry = block.get("geometry") or {}
        zone_id = str(geometry.get("zoneId") or "")
        bbox = _bbox_pt(geometry.get("bboxPx"), sx, sy)
        if bbox is None:
            raise RuntimeError(f"Canonical Word build blocked: missing bbox for {block_id}")

        if current_zone is not None and zone_id != current_zone:
            breaker = doc.add_paragraph()
            breaker.paragraph_format.space_before = Pt(0)
            breaker.paragraph_format.space_after = Pt(0)
            breaker.add_run().add_break(WD_BREAK.COLUMN)
            current_zone = zone_id

        prev_bottom = previous_bottom_by_zone.get(zone_id)
        raw_gap = max(0.0, bbox[1] - prev_bottom) if prev_bottom is not None else 0.0
        gap_pt = min(raw_gap, 12.0)
        size = _font_size_pt(block, sy)
        container_id = block_to_container.get(block_id)
        container = containers.get(container_id or "") if container_id else None

        if semantic == "figure":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(gap_pt)
            p.paragraph_format.space_after = Pt(0)
            asset = _resolve_asset(Path(package_root) if package_root else None, block)
            if asset is None:
                raise RuntimeError(
                    f"Canonical Word build blocked: figure asset unresolved for {block_id}"
                )
            width = max(12.0, bbox[2] - bbox[0])
            try:
                image_audit = _add_word_picture(p, asset, width)
            except Exception as exc:
                raise RuntimeError(
                    f"Canonical Word build blocked: figure asset failed {asset}: {exc}"
                ) from exc
            report_items.append({
                "id": block_id,
                "type": semantic,
                "status": "rendered",
                **image_audit,
            })
            previous_bottom_by_zone[zone_id] = bbox[3]
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(gap_pt)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.widow_control = True
        p.paragraph_format.keep_together = semantic in {"heading", "equation"}
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(max(size * 1.18, size + 1.0))

        if semantic == "heading":
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(max(size, 9.0))
            run.font.color.rgb = RGBColor(0, 0, 255)
            _set_language(run)
        elif semantic == "equation":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if not _append_omml(p, str(content.get("rawMarkdown") or text)):
                run = p.add_run(_latex_body(text))
                run.font.size = Pt(max(size, 8.0))
                _set_language(run)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(max(5.5, size))
            _set_language(run)

        frame_audit = _paragraph_frame_style(p, container) if container is not None else None
        report_items.append({
            "id": block_id,
            "type": semantic,
            "status": "rendered",
            "zoneId": zone_id,
            "bboxPt": bbox,
            "fontSizePt": size,
            "containerId": container_id,
            "containerStyle": frame_audit,
        })
        previous_bottom_by_zone[zone_id] = bbox[3]

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return {
        "version": VERSION,
        "status": "rendered",
        "targetPage": target_page,
        "output": str(output_path),
        "sourcePolicy": {
            "content": "canonical.content",
            "semantics": "canonical.semantic",
            "geometry": "canonical.geometry",
            "margins": "canonical.pageTopology.recoveredFrameEvidence",
            "zones": "canonical.pageTopology.zones",
            "crossZoneOrder": "canonical.pageTopology.crossZoneReadingOrder",
            "visualContainers": "canonical.pdfContainers",
            "matching": "forbidden",
            "reinterpretation": "forbidden",
            "legacyFallback": "forbidden",
        },
        "pageCanvasPx": [width_px, height_px],
        "recoveredFramePt": frame,
        "columns": columns_audit,
        "readingOrder": reading_order,
        "itemCount": len(report_items),
        "items": report_items,
    }


__all__ = ["VERSION", "build_canonical_native_document"]
