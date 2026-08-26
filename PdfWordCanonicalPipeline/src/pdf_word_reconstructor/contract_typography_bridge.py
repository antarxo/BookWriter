from __future__ import annotations

from contextlib import contextmanager
from typing import Any, Iterator

from docx.shared import Pt, RGBColor


def _font_name(typography: dict[str, Any]) -> str | None:
    value = ((typography.get("fontFamily") or {}).get("dominant"))
    return str(value).strip() if value else None


def _font_size(typography: dict[str, Any]) -> float | None:
    value = ((typography.get("fontSizePt") or {}).get("dominant"))
    try:
        return float(value) if value is not None else None
    except (TypeError, ValueError):
        return None


def _rgb_hex(value: Any) -> str | None:
    if isinstance(value, str):
        text = value.strip().lstrip("#")
        if len(text) == 6 and all(ch in "0123456789abcdefABCDEF" for ch in text):
            return text.upper()
        return None
    try:
        number = int(value)
    except (TypeError, ValueError):
        return None
    return f"{number & 0xFFFFFF:06X}"


def _pdf_color(value: Any) -> RGBColor | None:
    text = _rgb_hex(value)
    if not text:
        return None
    return RGBColor(int(text[0:2], 16), int(text[2:4], 16), int(text[4:6], 16))


def _contract_run(text: str, item: dict[str, Any]) -> dict[str, Any]:
    typography = item.get("__pdfTypography") if isinstance(item.get("__pdfTypography"), dict) else {}
    emphasis = typography.get("emphasis") if isinstance(typography.get("emphasis"), dict) else {}
    return {
        "text": text,
        "bold": bool(isinstance(emphasis.get("boldRatio"), (int, float)) and float(emphasis["boldRatio"]) >= 0.55),
        "italic": bool(isinstance(emphasis.get("italicRatio"), (int, float)) and float(emphasis["italicRatio"]) >= 0.55),
        "underline": bool(isinstance(emphasis.get("underlineRatio"), (int, float)) and float(emphasis["underlineRatio"]) >= 0.55),
        "superscript": bool(isinstance(emphasis.get("superscriptRatio"), (int, float)) and float(emphasis["superscriptRatio"]) >= 0.55),
        "font": _font_name(typography),
        "size_pt": _font_size(typography),
        "pdf_color": ((typography.get("color") or {}).get("dominant")),
        "__contract": True,
    }


def _contract_by_id(contract: dict[str, Any]) -> dict[str, dict[str, Any]]:
    return {
        str(item.get("id")): item
        for item in contract.get("items", []) or []
        if item.get("id")
    }


def _region_contracts(contract: dict[str, Any], page_structure: dict[str, Any]) -> dict[tuple[str, ...], dict[str, Any]]:
    by_slot: dict[tuple[int, str], dict[str, Any]] = {}
    for item in contract.get("items", []) or []:
        placement = item.get("placement") or {}
        page = int(placement.get("page") or 0)
        slot = str(placement.get("slotId") or "")
        if page and slot:
            by_slot[(page, slot)] = item
    result: dict[tuple[str, ...], dict[str, Any]] = {}
    for page in page_structure.get("pages", []) or []:
        page_no = int(page.get("page") or 0)
        for item in page.get("flow", []) or []:
            slot = str(item.get("id") or "")
            contract_item = by_slot.get((page_no, slot))
            if contract_item is None and slot.startswith("flow-"):
                contract_item = by_slot.get((page_no, slot[5:]))
            if contract_item is None:
                continue
            region_ids = tuple(str(value) for value in item.get("region_ids", []) or [])
            if region_ids:
                result[region_ids] = contract_item
    return result


def _apply_contract_to_structure(contract: dict[str, Any], page_structure: dict[str, Any]) -> None:
    by_slot: dict[tuple[int, str], dict[str, Any]] = {}
    for contract_item in contract.get("items", []) or []:
        placement = contract_item.get("placement") or {}
        page = int(placement.get("page") or 0)
        slot = str(placement.get("slotId") or "")
        if page and slot:
            by_slot[(page, slot)] = contract_item
    for page in page_structure.get("pages", []) or []:
        page_no = int(page.get("page") or 0)
        for collection_name in ("flow", "callouts"):
            for item in page.get(collection_name, []) or []:
                slot = str(item.get("id") or "")
                contract_item = by_slot.get((page_no, slot))
                if contract_item is None and slot.startswith("flow-"):
                    contract_item = by_slot.get((page_no, slot[5:]))
                if contract_item is None:
                    continue
                item["__wordParagraph"] = contract_item.get("wordParagraph") or {}
                item["__pdfTypography"] = contract_item.get("pdfTypography") or {}
                item["__buildContractId"] = contract_item.get("id")


def _alignment_value(legacy_module: Any, value: Any):
    return {
        "left": legacy_module.WD_ALIGN_PARAGRAPH.LEFT,
        "center": legacy_module.WD_ALIGN_PARAGRAPH.CENTER,
        "right": legacy_module.WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": legacy_module.WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get(str(value or "").strip().lower())


def _float_or_none(value: Any) -> float | None:
    try:
        return float(value) if value is not None else None
    except (TypeError, ValueError):
        return None


def _apply_paragraph_contract(legacy_module: Any, paragraph: Any, contract_item: dict[str, Any]) -> dict[str, Any]:
    word = contract_item.get("wordParagraph") if isinstance(contract_item.get("wordParagraph"), dict) else {}
    geometry = word.get("geometry") if isinstance(word.get("geometry"), dict) else {}
    spacing = word.get("spacing") if isinstance(word.get("spacing"), dict) else {}
    applied: dict[str, Any] = {"buildContractId": contract_item.get("id")}

    alignment = geometry.get("alignment") if isinstance(geometry.get("alignment"), dict) else {}
    align = _alignment_value(legacy_module, alignment.get("value"))
    if align is not None:
        paragraph.alignment = align
        applied["alignment"] = alignment.get("value")

    left = _float_or_none(geometry.get("leftIndentPt"))
    right = _float_or_none(geometry.get("rightIndentPt"))
    first = _float_or_none(geometry.get("firstLineIndentPt")) or 0.0
    hanging = _float_or_none(geometry.get("hangingIndentPt")) or 0.0
    if left is not None:
        paragraph.paragraph_format.left_indent = Pt(max(0.0, left))
        applied["leftIndentPt"] = round(max(0.0, left), 3)
    if right is not None:
        paragraph.paragraph_format.right_indent = Pt(max(0.0, right))
        applied["rightIndentPt"] = round(max(0.0, right), 3)
    paragraph.paragraph_format.first_line_indent = Pt(first - hanging)
    applied["firstLineIndentPt"] = round(first, 3)
    applied["hangingIndentPt"] = round(hanging, 3)

    line_height = _float_or_none(geometry.get("lineHeightPt"))
    if line_height is not None and line_height > 0:
        paragraph.paragraph_format.line_spacing_rule = legacy_module.WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(line_height)
        applied["lineHeightPt"] = round(line_height, 3)

    before = _float_or_none(spacing.get("spaceBeforePt"))
    after = _float_or_none(spacing.get("spaceAfterPt"))
    if before is not None:
        paragraph.paragraph_format.space_before = Pt(max(0.0, before))
        applied["spaceBeforePt"] = round(max(0.0, before), 3)
    if after is not None:
        paragraph.paragraph_format.space_after = Pt(max(0.0, after))
        applied["spaceAfterPt"] = round(max(0.0, after), 3)
    return applied


def _apply_contract_run(legacy_module: Any, paragraph: Any, text: str, item: dict[str, Any]) -> None:
    src = _contract_run(text, item)
    run = paragraph.add_run(text)
    run.bold = bool(src.get("bold"))
    run.italic = bool(src.get("italic"))
    run.underline = bool(src.get("underline"))
    run.font.superscript = bool(src.get("superscript"))
    font = str(src.get("font") or "").strip()
    if font:
        run.font.name = font
    size = src.get("size_pt")
    if isinstance(size, (int, float)):
        run.font.size = Pt(float(size))
    pdf_rgb = _pdf_color(src.get("pdf_color"))
    if pdf_rgb is not None:
        run.font.color.rgb = pdf_rgb
    legacy_module._set_run_language(run)


def _dash_style(value: Any) -> str:
    text = str(value or "").strip().lower()
    if not text or text in {"[] 0", "[]0", "none"}:
        return "single"
    return "dashed"


def _apply_frame_style(legacy_module: Any, paragraph: Any, frame: dict[str, Any]) -> dict[str, Any]:
    border = frame.get("border") if isinstance(frame.get("border"), dict) else {}
    fill = frame.get("fill") if isinstance(frame.get("fill"), dict) else {}
    applied: dict[str, Any] = {
        "evidenceStatus": frame.get("evidenceStatus"),
        "evidenceConfidence": frame.get("evidenceConfidence"),
        "drawingId": frame.get("drawingId"),
        "borderApplied": False,
        "fillApplied": False,
    }
    if frame.get("evidenceStatus") != "matched" or frame.get("evidenceConfidence") not in {"high", "medium"}:
        applied["reason"] = "no-confident-vector-style"
        return applied

    p_pr = paragraph._p.get_or_add_pPr()
    border_color = _rgb_hex(border.get("color"))
    border_opacity = _float_or_none(border.get("opacity"))
    border_width = _float_or_none(border.get("widthPt"))
    if border.get("status") == "extracted" and border_color and border_width is not None:
        if border_opacity is None or border_opacity >= 0.99:
            borders = p_pr.find(legacy_module.qn("w:pBdr"))
            if borders is None:
                borders = legacy_module.OxmlElement("w:pBdr")
                p_pr.append(borders)
            size_eighth_points = max(2, min(96, round(max(0.25, border_width) * 8.0)))
            for edge in ("top", "left", "bottom", "right"):
                node = borders.find(legacy_module.qn(f"w:{edge}"))
                if node is None:
                    node = legacy_module.OxmlElement(f"w:{edge}")
                    borders.append(node)
                node.set(legacy_module.qn("w:val"), _dash_style(border.get("dashes")))
                node.set(legacy_module.qn("w:sz"), str(size_eighth_points))
                node.set(legacy_module.qn("w:space"), "0")
                node.set(legacy_module.qn("w:color"), border_color)
            applied["borderApplied"] = True
            applied["borderColor"] = f"#{border_color}"
            applied["borderWidthPt"] = round(border_width, 3)
            applied["borderStyle"] = _dash_style(border.get("dashes"))
        else:
            applied["borderReason"] = "unsupported-nonopaque-pdf-stroke"
    else:
        applied["borderReason"] = border.get("status") or "no-painted-stroke"

    fill_color = _rgb_hex(fill.get("color"))
    fill_opacity = _float_or_none(fill.get("opacity"))
    if fill.get("status") == "extracted" and fill_color:
        if fill_opacity is None or fill_opacity >= 0.99:
            shd = p_pr.find(legacy_module.qn("w:shd"))
            if shd is None:
                shd = legacy_module.OxmlElement("w:shd")
                p_pr.append(shd)
            shd.set(legacy_module.qn("w:val"), "clear")
            shd.set(legacy_module.qn("w:color"), "auto")
            shd.set(legacy_module.qn("w:fill"), fill_color)
            applied["fillApplied"] = True
            applied["fillColor"] = f"#{fill_color}"
        else:
            applied["fillReason"] = "unsupported-nonopaque-pdf-fill"
    else:
        applied["fillReason"] = fill.get("status") or "no-painted-fill"
    return applied


@contextmanager
def contract_typography_bridge(legacy_module: Any, contract: dict[str, Any], page_structure: dict[str, Any]) -> Iterator[dict[str, Any]]:
    """Route Word rendering through the maps-first build contract."""
    _apply_contract_to_structure(contract, page_structure)
    region_map = _region_contracts(contract, page_structure)
    contract_by_id = _contract_by_id(contract)

    original_document = legacy_module.Document
    original_item_text = legacy_module._item_text_and_runs
    original_add_runs = legacy_module._add_runs
    original_dominant = legacy_module._dominant_span_size
    original_max = legacy_module._max_span_size
    original_line_height = legacy_module._line_height
    original_gap = legacy_module._word_flow_gap
    original_add_callout = legacy_module._add_callout

    pending: dict[str, Any] = {"contractItem": None}
    paragraph_bindings: list[tuple[Any, dict[str, Any]]] = []
    callout_rows: list[dict[str, Any]] = []
    audit: dict[str, Any] = {
        "policy": "build-contract-paragraph-format-and-vector-frame-before-save",
        "boundParagraphCount": 0,
        "appliedParagraphCount": 0,
        "calloutCount": 0,
        "callouts": callout_rows,
        "items": [],
    }

    def document_factory(*args, **kwargs):
        doc = original_document(*args, **kwargs)
        original_add_paragraph = doc.add_paragraph
        original_save = doc.save

        def add_paragraph(*p_args, **p_kwargs):
            paragraph = original_add_paragraph(*p_args, **p_kwargs)
            contract_item = pending.get("contractItem")
            if isinstance(contract_item, dict):
                paragraph_bindings.append((paragraph, contract_item))
                audit["boundParagraphCount"] = len(paragraph_bindings)
                pending["contractItem"] = None
            return paragraph

        def save(*s_args, **s_kwargs):
            applied_rows: list[dict[str, Any]] = []
            for paragraph, contract_item in paragraph_bindings:
                applied_rows.append(_apply_paragraph_contract(legacy_module, paragraph, contract_item))
            audit["appliedParagraphCount"] = len(applied_rows)
            audit["items"] = applied_rows
            return original_save(*s_args, **s_kwargs)

        doc.add_paragraph = add_paragraph
        doc.save = save
        return doc

    def item_text_and_runs(item: dict[str, Any], matches: dict[str, Any], docx_paras: dict[str, Any]):
        text = str(item.get("text") or "")
        if item.get("type") == "text" and item.get("__pdfTypography"):
            pending["contractItem"] = contract_by_id.get(str(item.get("__buildContractId") or ""))
            return text, [_contract_run(text, item)], "markdown-via-build-contract", []
        return original_item_text(item, matches, docx_paras)

    def add_runs(paragraph, text, source_runs, font_size, color=None, force_bold=False, italic=False, preserve_line_breaks=False):
        if source_runs and any(bool(run.get("__contract")) for run in source_runs):
            for src in source_runs:
                value = str(src.get("text") or "")
                if not value:
                    continue
                parts = value.split("\n") if preserve_line_breaks else [" ".join(value.split())]
                for index, part in enumerate(parts):
                    run = paragraph.add_run(part)
                    run.bold = bool(src.get("bold"))
                    run.italic = bool(src.get("italic"))
                    run.underline = bool(src.get("underline"))
                    run.font.superscript = bool(src.get("superscript"))
                    font = str(src.get("font") or "").strip()
                    if font:
                        run.font.name = font
                    size = src.get("size_pt")
                    run.font.size = Pt(float(size) if isinstance(size, (int, float)) else float(font_size))
                    pdf_rgb = _pdf_color(src.get("pdf_color"))
                    if pdf_rgb is not None:
                        run.font.color.rgb = pdf_rgb
                    legacy_module._set_run_language(run)
                    if preserve_line_breaks and index < len(parts) - 1:
                        run.add_break(legacy_module.WD_BREAK.LINE)
            return
        return original_add_runs(paragraph, text, source_runs, font_size, color, force_bold, italic, preserve_line_breaks)

    def add_callout(doc, callout, text, body_size, source_paragraphs=None, paragraph_ids=None):
        contract_item = contract_by_id.get(str(callout.get("__buildContractId") or ""))
        if not contract_item:
            raise RuntimeError(f"Maps-first callout build blocked: no contract for {callout.get('id')}")
        word = contract_item.get("wordParagraph") if isinstance(contract_item.get("wordParagraph"), dict) else {}
        frame = word.get("frame") if isinstance(word.get("frame"), dict) else {}
        typography = contract_item.get("pdfTypography") if isinstance(contract_item.get("pdfTypography"), dict) else {}
        bbox = frame.get("bboxPt") or callout.get("bbox")
        if not isinstance(bbox, (list, tuple)) or len(bbox) != 4:
            raise RuntimeError(f"Maps-first callout build blocked: missing frame bbox for {callout.get('id')}")
        size = _font_size(typography)
        line_height = _float_or_none(((word.get("geometry") or {}).get("lineHeightPt")))
        if size is None or line_height is None or line_height <= 0:
            raise RuntimeError(f"Maps-first callout build blocked: missing PDF typography/line-height for {callout.get('id')}")

        pending["contractItem"] = contract_item
        paragraph = doc.add_paragraph()
        legacy_module._set_frame(paragraph, list(map(float, bbox)))
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.widow_control = False
        _apply_contract_run(legacy_module, paragraph, str(text or ""), callout)
        frame_style = _apply_frame_style(legacy_module, paragraph, frame)

        border = frame.get("border") if isinstance(frame.get("border"), dict) else {}
        fill = frame.get("fill") if isinstance(frame.get("fill"), dict) else {}
        row = {
            "id": callout.get("id"),
            "buildContractId": contract_item.get("id"),
            "bbox": list(map(float, bbox)),
            "font_size_pt": round(size, 3),
            "line_height_pt": round(line_height, 3),
            "contentSource": "markdown-via-build-contract",
            "typographySource": "pdf-via-build-contract",
            "frameSource": frame.get("geometrySource") or "wordParagraph.frame",
            "drawingId": frame.get("drawingId"),
            "borderStatus": border.get("status") or "unresolved",
            "fillStatus": fill.get("status") or "unresolved",
            "frameStyle": frame_style,
            "inventedBorder": False,
            "inventedFill": False,
            "fontShrinkApplied": False,
            "native_math_count": 0,
            "source_paragraphs": [],
            "contained_visual_groups": list(callout.get("contained_visual_groups", [])),
        }
        callout_rows.append(row)
        audit["calloutCount"] = len(callout_rows)
        return row

    def contract_for_regions(region_ids: Any) -> dict[str, Any] | None:
        return region_map.get(tuple(str(value) for value in region_ids or []))

    def dominant(regions_by_id, region_ids, fallback):
        item = contract_for_regions(region_ids)
        if item:
            size = _font_size(item.get("pdfTypography") or {})
            if size is not None:
                return size
        return original_dominant(regions_by_id, region_ids, fallback)

    def maximum(regions_by_id, region_ids, fallback):
        item = contract_for_regions(region_ids)
        if item:
            profile = ((item.get("pdfTypography") or {}).get("fontSizePt") or {}).get("profile") or []
            values: list[float] = []
            for row in profile:
                try:
                    values.append(float(row.get("value")))
                except (TypeError, ValueError):
                    pass
            if values:
                return max(values)
            size = _font_size(item.get("pdfTypography") or {})
            if size is not None:
                return size
        return original_max(regions_by_id, region_ids, fallback)

    def line_height(regions_by_id, region_ids, fallback):
        item = contract_for_regions(region_ids)
        if item:
            value = (((item.get("wordParagraph") or {}).get("geometry") or {}).get("lineHeightPt"))
            try:
                if value is not None and float(value) > 0:
                    return float(value)
            except (TypeError, ValueError):
                pass
        return original_line_height(regions_by_id, region_ids, fallback)

    def flow_gap(raw_gap, item, body_size, gap_scale):
        spacing = ((item.get("__wordParagraph") or {}).get("spacing") or {})
        value = spacing.get("spaceBeforePt")
        try:
            if value is not None:
                gap = max(0.0, float(value))
                return {
                    "raw_gap_pt": round(gap, 2),
                    "scaled_gap_pt": round(gap, 2),
                    "applied_gap_pt": round(gap, 2),
                    "gap_clamped": False,
                    "gap_quantized": False,
                    "gap_policy": "build-contract-space-before",
                }
        except (TypeError, ValueError):
            pass
        return original_gap(raw_gap, item, body_size, gap_scale)

    legacy_module.Document = document_factory
    legacy_module._item_text_and_runs = item_text_and_runs
    legacy_module._add_runs = add_runs
    legacy_module._dominant_span_size = dominant
    legacy_module._max_span_size = maximum
    legacy_module._line_height = line_height
    legacy_module._word_flow_gap = flow_gap
    legacy_module._add_callout = add_callout
    try:
        yield audit
    finally:
        legacy_module.Document = original_document
        legacy_module._item_text_and_runs = original_item_text
        legacy_module._add_runs = original_add_runs
        legacy_module._dominant_span_size = original_dominant
        legacy_module._max_span_size = original_max
        legacy_module._line_height = original_line_height
        legacy_module._word_flow_gap = original_gap
        legacy_module._add_callout = original_add_callout
